"""Build the customer-facing SLIDE one-row and two-row formula reference."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "СЛАЙД_формулы_1_и_2_ряда.docx"

ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
TEXT = "1F2933"
MUTED = "5F6B76"
GRID = "AAB4BE"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F3EC"
LIGHT_YELLOW = "FFF4CC"
LIGHT_RED = "FCE8E6"
FORMULA_FILL = "F5F8FB"
WHITE = "FFFFFF"

TABLE_WIDTH_DXA = 10062
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str = TEXT,
    size: float = 9.2,
    font: str = "Calibri",
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, name=font, size=size, color=color, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    font_size: float = 9.1,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            color=ACCENT_DARK,
            size=font_size,
        )
        set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], str(value), size=font_size)
        prevent_row_split(table.rows[-1])

    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_status_box(doc: Document, label: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), label, bold=True, color=ACCENT_DARK, size=9.5)
    set_cell_text(table.cell(0, 1), text, size=9.5)
    set_cell_shading(table.cell(0, 0), fill)
    set_cell_shading(table.cell(0, 1), fill)
    set_table_geometry(table, [1900, 8162])
    prevent_row_split(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_formula(doc: Document, formula: str, note: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, FORMULA_FILL)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2 if note else 0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(formula)
    set_run_font(run, name="Consolas", size=9.2, color=ACCENT_DARK, bold=True)
    if note:
        note_paragraph = cell.add_paragraph()
        note_paragraph.paragraph_format.space_after = Pt(0)
        note_paragraph.paragraph_format.line_spacing = 1.05
        note_run = note_paragraph.add_run(note)
        set_run_font(note_run, size=8.9, color=MUTED, italic=True)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    prevent_row_split(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=10.6, color=TEXT)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    bold_prefix: str = "",
    color: str = TEXT,
    italic: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_run_font(prefix_run, size=10.8, color=color, bold=True)
        body_run = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(body_run, size=10.8, color=color, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.8, color=color, italic=italic)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, ACCENT, 18, 10),
        2: (13, ACCENT, 14, 7),
        3: (12, ACCENT_DARK, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(
        header.add_run("RALUMA | СЛАЙД | РАСЧЁТНЫЕ ФОРМУЛЫ"),
        size=8.5,
        color=MUTED,
        bold=True,
    )
    add_page_number(section.footer.paragraphs[0])


def add_title_page(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("СЛАЙД"), size=24, color=TEXT, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(
        subtitle.add_run("Формулы расчёта: стандарт 1 ряд и стандарт 2 ряда"),
        size=15,
        color=MUTED,
        bold=True,
    )

    metadata = [
        ("Проект", "RALUMA"),
        ("Назначение", "Согласование расчётной логики системы СЛАЙД"),
        ("Версия", "1.1"),
        ("Дата", date.today().strftime("%d.%m.%Y")),
        (
            "Статус",
            "Формулы сведены по действующему расчётному модулю и подтверждённым правилам.",
        ),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}: "), size=11, color=TEXT, bold=True)
        set_run_font(paragraph.add_run(value), size=11, color=TEXT)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT)
    p_bdr.append(bottom)
    rule._p.get_or_add_pPr().append(p_bdr)

    add_paragraph(
        doc,
        "Документ фиксирует единый порядок расчёта физических панелей, стекла, "
        "стекольного профиля RS2021, нарезки профилей, щёточных уплотнителей, "
        "фурнитуры, заглушек и крепежа. Схемы и группировка строк документов "
        "строятся после расчёта физических панелей и не изменяют левую/правую "
        "расчётную сторону.",
    )

    add_status_box(
        doc,
        "ВАЖНО",
        "В форме используются два независимых отступа под ручку: левый a и правый b. "
        "Сначала a и b нормализуются с учётом выбранных ручек, затем определяются "
        "krlp и krrp и рассчитывается базовая ширина промежуточного стекла. К левой "
        "и правой физической панели применяется только её собственный отступ.",
        LIGHT_GREEN,
    )
    add_status_box(
        doc,
        "ПОРЯДОК",
        "Физические панели → размеры стекла → RS2021 → группировка одинаковых "
        "позиций → производственный лист, заказ стекла и схемы.",
        LIGHT_BLUE,
    )

    add_heading(doc, "Состав документа", 2)
    for item in (
        "Обозначения и общие правила.",
        "Формулы СЛАЙД стандарт 1 ряд.",
        "Формулы СЛАЙД стандарт 2 ряда.",
        "Профили, RS2021, уплотнители, ролики и предупреждения.",
        "Фурнитура, заглушки, замки, ручки и крепёж.",
        "Окрашивание, нарезка и округления.",
        "Контрольный пример с независимыми отступами a и b.",
    ):
        add_bullet(doc, item)


def add_notation(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "1. Обозначения и общие правила", 1)
    add_table(
        doc,
        ["Переменная", "Смысл", "Единица / значение"],
        [
            ["W", "Габаритная ширина секции", "мм"],
            ["H", "Габаритная высота секции", "мм"],
            ["P", "Количество панелей", "шт"],
            ["Q", "Количество одинаковых секций", "шт"],
            ["a", "Отступ под ручку для левой физической панели", "мм"],
            ["b", "Отступ под ручку для правой физической панели", "мм"],
            ["c", "Отступ центральной ручки для системы 2 ряда", "мм"],
            ["O", "Перехлёст между стеклами", "0 / 9,5 / 11,5 мм"],
            ["M", "Базовая ширина промежуточного стекла", "мм"],
            ["L, R", "Ширина левого и правого крайнего стекла", "мм"],
            ["C", "Ширина центрального стекла в системе 2 ряда", "мм"],
            ["N", "Количество задействованных рельсов", "P / 2 для 2 рядов"],
        ],
        [1500, 6100, 2462],
    )

    add_heading(doc, "1.1. Перехлёст стекол", 2)
    add_table(
        doc,
        ["Межстекольный профиль", "Артикул", "O, мм"],
        [
            ["Алюминиевый", "RS2061", "9,5"],
            ["Прозрачный межстекольный", "RS1006", "9,5"],
            ["Профиль с зацепом", "RS3061", "11,5"],
            ["Без межстекольного профиля", "—", "0"],
        ],
        [4300, 2200, 3562],
    )

    add_heading(doc, "1.2. Боковые расчётные переменные", 2)
    add_table(
        doc,
        ["Переменная", "Левая сторона", "Правая сторона", "Условие"],
        [
            ["ppl / ppr", "16", "16", "Выбран пристеночный профиль"],
            ["rpl / rpr", "60", "60", "Выбран боковой профиль-замок RS2081"],
            [
                "rpl / rpr",
                "28",
                "28",
                "Выбран П-профиль RS1082 без пузырькового RS1002",
            ],
            ["pzl / pzr", "6", "6", "Выбран пузырьковый уплотнитель RS1002"],
            ["krlr / krrr", "8", "8", "Выбрана ручка-профиль RS112"],
            [
                "krlp / krrp",
                "16 при a = 0",
                "16 при b = 0",
                "Выбраны RS1082 и RS1002 на соответствующей стороне; иначе 0",
            ],
            ["pl / pr", "2", "2", "Выбраны RS1082 и RS1002 на одной стороне"],
            [
                "a / b",
                "вводится",
                "вводится",
                "Нормализуются до расчёта krlp / krrp",
            ],
        ],
        [1600, 1650, 1650, 5162],
        font_size=8.8,
    )

    add_status_box(
        doc,
        "НОРМАЛИЗАЦИЯ",
        "Сохранённый старый отступ у ручки, которая не поддерживает отступ, "
        "игнорируется: соответствующее a или b принимается равным 0. Только после "
        "этого определяется компенсация 16 мм для RS1082 + RS1002.",
        LIGHT_BLUE,
    )

    add_status_box(
        doc,
        "СТОРОНЫ",
        "Левая физическая панель использует только левые параметры; правая — только "
        "правые. Направление движения, нумерация и визуальный порядок на схеме не "
        "переносят поправки между сторонами.",
        LIGHT_YELLOW,
    )


def add_common_geometry(doc: Document) -> None:
    add_heading(doc, "2. Общая геометрия профилей", 1)
    add_heading(doc, "2.1. Горизонтальные и пристеночные профили", 2)
    add_formula(
        doc,
        "Lгор = W - 16 × Kстен",
        "Kстен — количество выбранных пристеночных профилей: 0, 1 или 2.",
    )
    add_table(
        doc,
        ["Группа", "3-рельсовая", "5-рельсовая", "Длина", "Количество"],
        [
            ["Верхняя направляющая", "RS1313", "RS1315", "Lгор", "Q"],
            ["Стандартный порог", "RS2323", "RS2325", "Lгор", "Q"],
            ["Накладной порог", "RS23231", "RS23251", "Lгор", "Q"],
            ["Пристеночный профиль", "RS2333", "RS2335", "H", "Kстен × Q"],
        ],
        [2700, 1700, 1700, 1900, 2062],
        font_size=8.8,
    )

    add_heading(doc, "2.2. Вертикальные размеры", 2)
    add_table(
        doc,
        ["Позиция", "Стандартный порог", "Накладной порог"],
        [
            ["Межстекольный профиль", "H - 162", "H - 150"],
            ["Боковой профиль-замок RS2081", "H - 65", "H - 55"],
            ["Ручка-профиль RS112", "H - 162", "H - 150"],
            ["Боковой П-профиль RS1082", "H - 65", "H - 55"],
            ["Высота стекла", "H - 106", "H - 94"],
            ["Пузырьковый уплотнитель RS1002", "Hстекла - 17", "Hстекла - 17"],
        ],
        [3650, 3206, 3206],
    )

    add_heading(doc, "2.3. Нарезка профиля длиннее 5950 мм", 2)
    add_formula(doc, "Lнорм = ceil(L)")
    add_formula(
        doc,
        "Если Lнорм ≤ 5950: одна часть Lнорм",
        "Предельная длина одной заготовки — 5950 мм.",
    )
    add_formula(doc, "Kчастей = ceil(Lнорм / 5950)")
    add_formula(
        doc,
        "Lбаза = floor(Lнорм / Kчастей); остаток = Lнорм - Lбаза × Kчастей",
        "Остаток распределяется по 1 мм между последними частями. Ни одна часть не "
        "превышает 5950 мм.",
    )


def add_one_row(doc: Document) -> None:
    add_heading(doc, "3. СЛАЙД стандарт 1 ряд", 1)
    add_status_box(
        doc,
        "АЛГОРИТМ",
        "Сначала нормализуются a и b, затем для каждой стороны определяется "
        "условная компенсация 16 мм и рассчитывается M — базовая ширина "
        "промежуточного стекла. После этого рассчитываются L и R. Группировка "
        "«Крайние / Промежуточные» выполняется только в конце.",
        LIGHT_GREEN,
    )

    add_heading(doc, "3.1. Ширина стекол", 2)
    add_formula(
        doc,
        "B1 = W - ppr - ppl - rpr - rpl - pzl - pzr"
        " - krlr - krlp - krrr - krrp - pl - pr - a - b"
        " + O × (P - 1)",
    )
    add_formula(doc, "M = B1 / P", "M — исходная ширина промежуточного стекла.")
    add_formula(
        doc,
        "L = M + a + krlr + krlp;   R = M + b + krrr + krrp",
        "a относится только к левой панели, b — только к правой. krlp = 16 только "
        "при a = 0; krrp = 16 только при b = 0. Стороны независимы.",
    )
    add_formula(
        doc,
        "RS1082 + RS1002: при a/b = 0 крайнее стекло = M + 16 мм",
        "Если на этой стороне задан ручной отступ, компенсация 16 мм равна 0 и к "
        "крайнему стеклу добавляется только ручной отступ этой стороны.",
    )
    add_formula(
        doc,
        "P = 1: M = W - ppr - ppl - pzl - pzr",
        "Однопанельная секция не использует формулу распределения P панелей.",
    )

    add_heading(doc, "3.2. Физические панели и группировка", 2)
    add_table(
        doc,
        ["Позиция", "Ширина", "Количество"],
        [
            ["Левая крайняя", "L", "Q"],
            ["Промежуточные", "M", "(P - 2) × Q, если P > 2"],
            ["Правая крайняя", "R", "Q"],
        ],
        [3200, 3200, 3662],
    )
    add_bullet(
        doc,
        "Если одинаковые по названию позиции имеют разную ширину или разный RS2021, "
        "они не объединяются до потери физической стороны.",
    )
    add_bullet(
        doc,
        "Параметр «1-я панель изнутри» влияет на номера, стрелки и тип RS107L/RS107R, "
        "но не меняет L и R местами в расчёте.",
    )

    add_heading(doc, "3.3. Количество профилей", 2)
    add_table(
        doc,
        ["Позиция", "Количество"],
        [
            ["Порог", "Q"],
            ["Верхняя направляющая", "Q"],
            ["Пристеночный профиль", "Kстен × Q"],
            ["Межстекольный профиль", "(P - 1) × Q"],
            ["RS2081", "Количество выбранных сторон × Q"],
            ["RS112", "Количество выбранных сторон × Q"],
            ["RS1082", "Количество выбранных сторон × Q"],
            ["RS1002", "Количество выбранных сторон × Q"],
            ["Защёлка в пол", "Количество выбранных сторон × Q"],
            ["RS2021", "По каждой физической панели, затем группировка по длине"],
        ],
        [5900, 4162],
    )

    add_heading(doc, "3.4. Щёточный уплотнитель", 2)
    add_formula(
        doc,
        "RU008 7×6 = (Lгор / 1000) × P × 2 × Q"
        " + (L_RS112 / 1000 + 0,03) × K_RS112 × Q",
    )
    add_formula(
        doc,
        "RU007 7×12 = (Lмеж / 1000 + 0,03) × (P - 1) × Q",
        "RU007 рассчитывается для межстекольных RS2061 и RS1006.",
    )


def add_two_row(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "4. СЛАЙД стандарт 2 ряда", 1)
    add_status_box(
        doc,
        "ПАНЕЛИ",
        "Для 2 рядов используются чётные количества панелей 4 / 6 / 8 / 10. "
        "Физический расчёт идёт слева направо; центральные левая и правая панели "
        "рассчитываются отдельно.",
        LIGHT_BLUE,
    )

    add_heading(doc, "4.1. Центральные переменные", 2)
    add_table(
        doc,
        ["Переменная", "Значение", "Условие"],
        [
            ["centr1", "46,5", "В центре выбраны две ручки-профиля RS112"],
            ["centr2", "8", "В центре выбраны две ручки-профиля RS112"],
            ["c", "вводится / 0", "Отступ центральной ручки для центральной фурнитуры"],
            ["Kцентр_RS112", "2", "Две центральные ручки-профиля"],
        ],
        [2200, 2200, 5662],
    )

    add_heading(doc, "4.2. Ширина стекол", 2)
    add_formula(
        doc,
        "B2 = W - 3 - ppr - ppl - rpr - rpl - pzl - pzr"
        " - krlr - krlp - krrr - krrp - pl - pr"
        " - centr1 - centr2 - a - b - 2c"
        " + O × (P - 2)",
    )
    add_formula(doc, "M = B2 / P")
    add_formula(
        doc,
        "L = M + a + krlr + krlp;   R = M + b + krrr + krrp",
        "krlp и krrp уже рассчитаны после нормализации a и b; каждая сторона "
        "использует только собственный ручной отступ или собственную компенсацию.",
    )
    add_formula(
        doc,
        "C = M + c + centr2",
        "C — ширина центрального стекла. Центральная левая и центральная правая "
        "панели остаются отдельными физическими позициями.",
    )

    add_heading(doc, "4.3. Физические панели", 2)
    add_table(
        doc,
        ["Позиция", "Ширина", "Количество"],
        [
            ["Левая крайняя", "L", "Q"],
            ["Левые промежуточные", "M", "(P - 4) / 2 × Q"],
            ["Центральная левая", "C", "Q"],
            ["Центральная правая", "C", "Q"],
            ["Правые промежуточные", "M", "(P - 4) / 2 × Q"],
            ["Правая крайняя", "R", "Q"],
        ],
        [3200, 3200, 3662],
    )

    add_heading(doc, "4.4. Количество профилей", 2)
    add_table(
        doc,
        ["Позиция", "Количество / длина"],
        [
            ["Межстекольный профиль", "(P - 2) × Q"],
            ["Боковые RS2081 / RS1082 / RS1002", "Количество выбранных сторон × Q"],
            [
                "RS112",
                "(Количество выбранных боковых сторон + 2 центральные) × Q",
            ],
            ["RS1083", "(L_RS112 + 17 мм), Q"],
            ["Магнитный уплотнитель RU010", "Длина RS1083, количество 2Q"],
            ["RS3110", "Lмеж, Q — когда центральный стык без пары RS112"],
            ["RS2021", "По каждой физической панели, затем группировка по длине"],
        ],
        [4400, 5662],
    )

    add_heading(doc, "4.5. Щёточный уплотнитель", 2)
    add_formula(doc, "N = P / 2", "N — количество задействованных рельсов.")
    add_formula(
        doc,
        "RU008 7×6 = (Lгор / 1000) × N × 2 × Q"
        " + (L_RS112 / 1000 + 0,03) × K_RS112_всего × Q",
    )
    add_formula(
        doc,
        "RU007 7×12 = (Lмеж / 1000 + 0,03) × (P - 2) × Q",
        "RU007 рассчитывается для межстекольных RS2061 и RS1006.",
    )


def add_glass_profile(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "5. Стекольный профиль RS2021", 1)
    add_status_box(
        doc,
        "ИСТОЧНИК",
        "RS2021 рассчитывается по каждой физической панели. Нельзя сначала "
        "объединять строки «Крайние» или «Центральные», потому что одинаковые "
        "стекла могут иметь разные длины RS2021.",
        LIGHT_GREEN,
    )

    add_heading(doc, "5.1. Боковые и промежуточные панели", 2)
    add_table(
        doc,
        ["Условие физической панели", "Длина RS2021"],
        [
            ["Глухая / неподвижная панель", "Ширина стекла"],
            ["На этой стороне выбрана ручка-профиль RS112", "Ширина стекла + 16 мм"],
            [
                "На этой стороне выбран RS1002, панель подвижная",
                "Ширина стекла - 3 мм",
            ],
            [
                "Промежуточная панель с межстекольным профилем",
                "Ширина стекла - 3 мм",
            ],
            ["Иное состояние", "Ширина стекла"],
        ],
        [6200, 3862],
    )

    add_heading(doc, "5.2. Центральные панели 2 рядов", 2)
    add_table(
        doc,
        ["Панель", "Условие", "Длина RS2021"],
        [
            [
                "Центральная левая",
                "Пара RS112, панель с соединительным профилем RS1083",
                "Стекло + 19 мм",
            ],
            [
                "Центральная правая",
                "Пара RS112, панель без RS1083",
                "Стекло + 16 мм",
            ],
            ["Центральная без RS112", "Есть межстекольный профиль", "Стекло - 3 мм"],
            ["Центральная без RS112", "Нет межстекольного профиля", "Стекло"],
        ],
        [2200, 5000, 2862],
        font_size=8.8,
    )

    add_heading(doc, "5.3. Группировка RS2021", 2)
    add_formula(doc, "L_RS2021_нарезка = round(L_RS2021_физической_панели)")
    add_formula(
        doc,
        "Ключ группировки = округлённая длина RS2021",
        "Объединяются только равные длины. Количество суммируется. Разные длины "
        "остаются отдельными строками, даже если стекла входят в одну группу.",
    )


def add_hardware(doc: Document) -> None:
    add_heading(doc, "6. Фурнитура и заглушки", 1)
    add_heading(doc, "6.1. Общие позиции", 2)
    add_table(
        doc,
        ["Артикул", "Назначение", "1 ряд", "2 ряда"],
        [
            ["RSD1", "Демпфер", "2(P - 1)Q", "2(P - 2)Q"],
            ["RSD2", "Компенсатор", "2(P - 1)Q", "2(P - 2)Q"],
            ["RS1121", "Накладка на RS112", "K_RS112 × Q", "K_RS112_всего × Q"],
            ["RS3018", "Замок-защёлка 1-стор.", "Выбранные стороны × Q", "Выбранные стороны × Q"],
            ["RS3020", "Замок двухсторонний", "Выбранные стороны × Q", "Выбранные стороны × Q"],
            ["RS122", "Ответная планка защёлки RS3018", "Количество RS3018", "Количество RS3018"],
            ["RS123", "Ответная планка замка RS3020", "Количество RS3020", "Количество RS3020"],
            ["RS206", "Накидная защёлка в центре", "—", "Выбор × Q"],
            ["RS30301", "Центральный замок", "—", "Выбор × Q"],
            ["RU1039", "Наклейка RALUMA", "Q", "Q"],
            ["RS150", "Инструкция СЛАЙД", "Q", "Q"],
        ],
        [1500, 4360, 2101, 2101],
        font_size=8.3,
    )

    add_heading(doc, "6.2. Ролики", 2)
    add_table(
        doc,
        ["Ширина физической панели", "Артикул", "Количество"],
        [
            ["Менее 350 мм", "—", "Предупреждение: уменьшите количество панелей"],
            ["От 350 до 500 мм включительно", "RU003, 2-колёсный", "2 × число панелей × Q"],
            ["Более 500 мм", "RU005, 4-колёсный", "2 × число панелей × Q"],
        ],
        [3400, 3300, 3362],
    )

    add_heading(doc, "6.3. Заглушки СЛАЙД 1 ряд", 2)
    add_table(
        doc,
        ["Артикул", "Количество"],
        [
            ["RS105, внутренние", "2(P - 1)Q"],
            [
                "RS106, крайние",
                "(левая неглухая + правая неглухая) × Q",
            ],
            ["RS107, запорные", "RS105 + RS106"],
            [
                "RS107L или RS107R",
                "(P - 1)Q для RS2061/RS1006; сторона по «1-й панели изнутри»",
            ],
        ],
        [5200, 4862],
    )

    add_heading(doc, "6.4. Заглушки СЛАЙД 2 ряда", 2)
    add_table(
        doc,
        ["Артикул", "Количество"],
        [
            ["RS105, внутренние", "2(P - 2)Q"],
            [
                "RS106, крайние",
                "(левая неглухая + правая неглухая) × Q",
            ],
            ["RS108, центральные", "2Q"],
            ["RS107, запорные", "RS105 + RS106; на RS108 не устанавливаются"],
            ["RS107L", "(P / 2 - 1)Q для RS2061/RS1006"],
            ["RS107R", "(P / 2 - 1)Q для RS2061/RS1006"],
        ],
        [5200, 4862],
    )


def add_fasteners(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "7. Крепёж", 1)
    add_formula(
        doc,
        "K_RS2081_на_сторону = max(8; round((H - 200) / 300))",
        "Одна линия крепления примерно на каждые 300 мм высоты, минимум 8 шт.",
    )
    add_table(
        doc,
        ["Крепёж", "Формула количества", "Комментарий"],
        [
            [
                "4,8×25 A2 DIN7982",
                "2 × (RS105 + RS106) — 1 ряд",
                "Прикрутить заглушки",
            ],
            [
                "4,8×25 A2 DIN7982",
                "2 × (RS105 + RS106 + RS108) — 2 ряда",
                "Прикрутить заглушки",
            ],
            [
                "3,9×13 A2 DIN7504M",
                "2 × (RU003 + RU005) + K_RS2081_на_сторону × Kсторон_RS2081 × Q",
                "Ролики и крепление RS2081",
            ],
            [
                "3,9×13 A2 DIN7504O",
                "7 × Kсторон_RS1082 × Q",
                "Крепление RS1082",
            ],
            [
                "3,5×13 A2 DIN7982",
                "2 × (RS122 + RS123)",
                "Прикрутить ответные планки RS122/123",
            ],
            [
                "5,4×25 A2 DIN912SW",
                "Kглухих_панелей × Q",
                "Крепление глухой панели",
            ],
        ],
        [2200, 5200, 2662],
        font_size=8.4,
    )

    add_heading(doc, "7.1. Саморез 4,8×38 A2 DIN7982 для порога", 2)
    add_table(
        doc,
        ["Порог", "Количество"],
        [
            ["3-рельсовый стандартный", "8"],
            ["5-рельсовый стандартный", "12"],
            ["3-рельсовый накладной", "4"],
            ["5-рельсовый накладной", "6"],
        ],
        [6500, 3562],
    )


def add_paint_and_rounding(doc: Document) -> None:
    add_heading(doc, "8. Окрашивание и документы", 1)
    add_bullet(
        doc,
        "Секция считается окрашиваемой, если тип окрашивания — RAL стандарт или RAL нестандарт.",
    )
    add_bullet(
        doc,
        "Анодированный порог не включается в заявку на покраску. Окрашенный "
        "стандартный или накладной порог включается.",
    )
    add_bullet(
        doc,
        "RS2061 включается в покраску; RS1006, RS3061, RS1002, RU010 и RS3110 не окрашиваются.",
    )
    add_bullet(
        doc,
        "Заявка на покраску группируется по цветам. Для каждого цвета формируется "
        "отдельная страница/группа.",
    )
    add_formula(
        doc,
        "Размер с припуском = ceil(чистовой размер / 50) × 50",
        "Округление вверх до ближайших 50 мм; ровное значение, например 3000, не меняется.",
    )
    add_formula(
        doc,
        "Общее, м.п. = количество × размер с припуском / 1000",
    )

    add_heading(doc, "8.1. Округления отображения", 2)
    add_table(
        doc,
        ["Значение", "Правило"],
        [
            ["Размеры стекла в расчёте", "Хранятся с точностью до 0,1 мм"],
            ["Размеры стекла в ПЛ/заказе", "Округление до целого миллиметра"],
            ["RS2021 в нарезке", "До ближайшего целого миллиметра"],
            ["Профиль перед распилом", "Вверх до целого миллиметра"],
            ["RU008 / RU007", "Вверх до 0,1 м; пример: 14,2 м и 3,4 м"],
            ["Площадь стекла", "Ширина × высота × количество / 1 000 000"],
        ],
        [3800, 6262],
    )
    add_bullet(
        doc,
        "Название и цвет стекла изменяют подписи, цвет схем и группировку заказа "
        "стекла, но сами геометрические формулы СЛАЙД не меняют.",
    )


def add_control_example(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "9. Контроль независимых отступов a и b", 1)
    add_status_box(
        doc,
        "ПОДТВЕРЖДЕНО",
        "В форме должны быть два поля: левый отступ a и правый отступ b. Один общий "
        "отступ использовать нельзя.",
        LIGHT_GREEN,
    )

    add_heading(doc, "9.1. Порядок расчёта", 2)
    steps = [
        "Нормализовать a и b с учётом выбранных левой и правой ручек.",
        "Определить krlp и krrp: 16 мм только для RS1082 + RS1002 при нулевом отступе соответствующей стороны, иначе 0.",
        "Определить остальные вычеты профилей отдельно для левой и правой стороны.",
        "Вычесть из общей ширины нормализованные a и b и рассчитанные krlp и krrp.",
        "Рассчитать базовую ширину промежуточного стекла M.",
        "Рассчитать L и R от M, прибавив только параметры соответствующей физической стороны.",
        "Рассчитать RS2021 отдельно по каждой физической панели.",
        "Только после этого объединять одинаковые строки для документов.",
    ]
    for index, text in enumerate(steps, start=1):
        add_paragraph(doc, f"{index}. {text}")

    add_heading(doc, "9.2. Эталон В26-5-4186", 2)
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Количество панелей", "3"],
            ["Правый отступ b", "100 мм"],
            ["krlp / krrp", "16 мм / 0 мм"],
            ["Промежуточное стекло M", "943,7 мм → 944 мм в документе"],
            ["Правое стекло R", "1043,7 мм → 1044 мм; ровно M + 100"],
            ["Левое крайнее стекло L", "959,7 мм → 960 мм; M + 16"],
            ["RS2021 слева", "960 мм"],
            ["RS2021 промежуточный", "941 мм"],
            ["RS2021 справа", "1041 мм"],
        ],
        [5200, 4862],
    )
    add_status_box(
        doc,
        "ПРОВЕРКА",
        "Изменение b должно менять правое стекло и базу распределения, но не должно "
        "переносить правую поправку на левую панель. То же правило симметрично для a.",
        LIGHT_YELLOW,
    )

    add_heading(doc, "9.3. Инварианты расчёта", 2)
    for item in (
        "Визуальный порядок панелей не меняет физические L и R.",
        "Первая панель изнутри влияет на схему и направление, но не переносит a и b.",
        "Глухая панель и подвижная панель определяются до вычисления RS2021.",
        "Стекло, RS2021, производственный лист и заказ стекла используют один массив физических панелей.",
        "Старые эталоны, где при ненулевом a или b всё ещё вычиталась компенсация 16 мм этой стороны, подлежат обновлению и не используются как контроль абсолютных размеров.",
    ):
        add_bullet(doc, item)


def add_summary(doc: Document) -> None:
    add_heading(doc, "10. Краткая карта зависимостей", 1)
    add_table(
        doc,
        ["Вход", "Что изменяет"],
        [
            ["W, H", "Стекла, вертикальные и горизонтальные профили, уплотнители"],
            ["P", "Количество стекол, стыков, профилей, роликов, заглушек и крепежа"],
            ["Q", "Все количественные позиции секции"],
            ["3 / 5 рельсов", "Порог, верхнюю направляющую, пристеночный профиль"],
            ["Стандартный / накладной порог", "Вертикальные вычеты и артикул порога"],
            ["RS2061 / RS1006 / RS3061", "Перехлёст O, RS2021, RU007 и окрашивание"],
            ["Левые профили и a", "Только левую физическую сторону"],
            ["Правые профили и b", "Только правую физическую сторону"],
            ["Центральные RS112 и c", "Только центральные панели 2 рядов"],
            ["Ширина каждой панели", "Тип ролика RU003 или RU005"],
            ["Глухая / подвижная", "RS2021, RS106, DIN912SW"],
            ["RAL / анод", "Состав заявки на покраску"],
        ],
        [3600, 6462],
        font_size=8.8,
    )
    add_status_box(
        doc,
        "ГЛАВНОЕ",
        "Расчёт начинается с физических панелей. Схемы, названия групп и порядок "
        "вывода являются представлением результата и не участвуют в формулах.",
        LIGHT_BLUE,
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_notation(doc)
    add_common_geometry(doc)
    add_one_row(doc)
    add_two_row(doc)
    add_glass_profile(doc)
    add_hardware(doc)
    add_fasteners(doc)
    add_paint_and_rounding(doc)
    add_control_example(doc)
    add_summary(doc)
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
