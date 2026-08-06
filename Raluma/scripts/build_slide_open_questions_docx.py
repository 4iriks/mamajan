"""Build the compact SLIDE pricing/quote RFI questionnaire."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SLIDE_открытые_вопросы_по_стоимости_и_КП.docx"

BLUE = "2E74B5"
NAVY = "163A4A"
MUTED = "5E6A73"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F3EC"
WHITE = "FFFFFF"
GRID = "AAB3BA"

TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(
    run,
    *,
    size: float = 11,
    color: str = "000000",
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
    keep_with_next: bool = False,
    alignment=None,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if alignment is not None:
        paragraph.alignment = alignment


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in CELL_MARGINS.items():
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width_node = properties.find(qn("w:tcW"))
    if width_node is None:
        width_node = OxmlElement("w:tcW")
        properties.append(width_node)
    width_node.set(qn("w:w"), str(width))
    width_node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: tuple[int, ...]) -> None:
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"Table widths must total {TABLE_WIDTH}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(TABLE_WIDTH))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT))
    table_indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)

    for row in table.rows:
        cannot_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cannot_split)
        for index, cell in enumerate(row.cells):
            set_cell_margins(cell)
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str = "000000",
    size: float = 9.5,
    alignment=None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph(
        paragraph,
        after=0,
        line=1.05,
        alignment=alignment,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run = OxmlElement("w:r")
        run.append(node)
        paragraph._p.append(run)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    set_paragraph(header_p, after=0, line=1.0)
    tab_stops = header_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))
    left = header_p.add_run("RALUMA | RFI")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_p.add_run("\tSLIDE · СТОИМОСТЬ И КП")
    set_run_font(right, size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(footer_p, after=0, line=1.0)
    run = footer_p.add_run("Страница ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer_p, "PAGE")
    run = footer_p.add_run(" из ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer_p, "NUMPAGES")


def add_masthead(document: Document) -> None:
    title = document.add_paragraph()
    set_paragraph(title, after=4, line=1.0, keep_with_next=True)
    run = title.add_run("RFI — ОТКРЫТЫЕ ВОПРОСЫ")
    set_run_font(run, size=23, bold=True)

    subtitle = document.add_paragraph()
    set_paragraph(subtitle, after=14, line=1.05, keep_with_next=True)
    run = subtitle.add_run("Стоимость и коммерческое предложение Raluma SLIDE")
    set_run_font(run, size=14, color="373737")

    rows = (
        ("Кому", "Руководителям продаж, финансов и производства"),
        ("От", "Рабочей группы SLIDE"),
        ("Дата", date(2026, 8, 6).strftime("%d.%m.%Y")),
        ("Тема", "Решения, необходимые для завершения первой версии ценообразования и КП"),
        ("Статус", "Требуется ответ по 8 вопросам"),
    )
    for label, value in rows:
        paragraph = document.add_paragraph()
        set_paragraph(paragraph, after=2, line=1.0, keep_with_next=True)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run)


def add_approved_callout(document: Document) -> None:
    table = document.add_table(rows=2, cols=1)
    set_table_geometry(table, (TABLE_WIDTH,))
    shade_cell(table.cell(0, 0), NAVY)
    set_cell_text(
        table.cell(0, 0),
        "УЖЕ УТВЕРЖДЕНО — ОТВЕТЫ НЕ ТРЕБУЮТСЯ",
        bold=True,
        color=WHITE,
        size=10,
    )
    shade_cell(table.cell(1, 0), PALE_GREEN)
    set_cell_text(
        table.cell(1, 0),
        (
            "Закалённое стекло с полными наименованиями; Excel-накладная; "
            "эскизы каждой секции в КП; публичные продажные цены профилей и "
            "фурнитуры без себестоимости, маржи и внутренних коэффициентов."
        ),
        size=10,
    )
    spacer = document.add_paragraph()
    set_paragraph(spacer, after=4, line=1.0)
    instruction = document.add_paragraph()
    set_paragraph(instruction, after=8, line=1.10)
    run = instruction.add_run(
        "Инструкция: отметьте один вариант в каждом вопросе, заполните поля "
        "«Уточнение» и при необходимости добавьте комментарий."
    )
    set_run_font(run, size=10.5, color=MUTED, italic=True)


def add_question(
    document: Document,
    number: int,
    question: str,
    options: tuple[tuple[str, str], ...],
    *,
    comment_prompt: str = "Комментарий / ограничения / ответственный: ________________________________________________",
) -> None:
    table = document.add_table(rows=3 + len(options), cols=3)
    widths = (6460, 840, 2060)
    set_table_geometry(table, widths)

    title = table.cell(0, 0).merge(table.cell(0, 2))
    set_cell_width(title, TABLE_WIDTH)
    shade_cell(title, NAVY)
    set_cell_text(
        title,
        f"{number}. {question}",
        bold=True,
        color=WHITE,
        size=10,
    )

    headers = ("Вариант решения", "Выбор", "Уточнение")
    for index, header in enumerate(headers):
        shade_cell(table.cell(1, index), LIGHT)
        set_cell_text(
            table.cell(1, index),
            header,
            bold=True,
            size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
    repeat_table_header(table.rows[1])

    for row_index, (option, clarification) in enumerate(options, start=2):
        set_cell_text(table.cell(row_index, 0), option, size=9.2)
        set_cell_text(
            table.cell(row_index, 1),
            "☐",
            size=13,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(table.cell(row_index, 2), clarification, size=9)

    comment = table.cell(len(options) + 2, 0).merge(
        table.cell(len(options) + 2, 2)
    )
    set_cell_width(comment, TABLE_WIDTH)
    shade_cell(comment, PALE_BLUE)
    set_cell_text(comment, comment_prompt, size=9, color=MUTED)

    spacer = document.add_paragraph()
    set_paragraph(spacer, after=5, line=1.0)


def add_defaults_question(document: Document) -> None:
    table = document.add_table(rows=7, cols=4)
    widths = (2500, 3660, 1100, 2100)
    set_table_geometry(table, widths)
    title = table.cell(0, 0).merge(table.cell(0, 3))
    set_cell_width(title, TABLE_WIDTH)
    shade_cell(title, NAVY)
    set_cell_text(
        title,
        "7. Каковы начальные значения НДС, срока действия, изготовления и оплаты?",
        bold=True,
        color=WHITE,
        size=10,
    )
    headers = ("Параметр", "Предлагаемое значение", "Принять", "Новое значение")
    for index, header in enumerate(headers):
        shade_cell(table.cell(1, index), LIGHT)
        set_cell_text(
            table.cell(1, index),
            header,
            bold=True,
            size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
    repeat_table_header(table.rows[1])
    values = (
        ("НДС", "Без НДС; ставка 20% сохранена для переключения"),
        ("Срок действия КП", "14 календарных дней"),
        ("Срок изготовления", "По согласованию"),
        ("Условия оплаты", "По согласованию"),
    )
    for row_index, (label, proposed) in enumerate(values, start=2):
        set_cell_text(table.cell(row_index, 0), label, bold=True, size=9.2)
        set_cell_text(table.cell(row_index, 1), proposed, size=9.2)
        set_cell_text(
            table.cell(row_index, 2),
            "☐",
            size=13,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(table.cell(row_index, 3), "________________", size=9)
    comment = table.cell(6, 0).merge(table.cell(6, 3))
    set_cell_width(comment, TABLE_WIDTH)
    shade_cell(comment, PALE_BLUE)
    set_cell_text(
        comment,
        "Комментарий (например, разные значения по ролям/клиентам): ________________________________________",
        size=9,
        color=MUTED,
    )
    spacer = document.add_paragraph()
    set_paragraph(spacer, after=5, line=1.0)


def add_scope_question(document: Document) -> None:
    table = document.add_table(rows=8, cols=4)
    widths = (3220, 1260, 1260, 3620)
    set_table_geometry(table, widths)
    title = table.cell(0, 0).merge(table.cell(0, 3))
    set_cell_width(title, TABLE_WIDTH)
    shade_cell(title, NAVY)
    set_cell_text(
        title,
        "8. Нужны ли в первой версии гарантия, условия доставки/монтажа, реквизиты и подписи?",
        bold=True,
        color=WHITE,
        size=10,
    )
    headers = ("Элемент КП", "Включить", "Отложить", "Уточнение / формулировка")
    for index, header in enumerate(headers):
        shade_cell(table.cell(1, index), LIGHT)
        set_cell_text(
            table.cell(1, index),
            header,
            bold=True,
            size=9,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
    repeat_table_header(table.rows[1])
    for row_index, label in enumerate(
        ("Гарантия", "Условия доставки", "Условия монтажа", "Реквизиты сторон", "Подписи"),
        start=2,
    ):
        set_cell_text(table.cell(row_index, 0), label, bold=True, size=9.2)
        for column in (1, 2):
            set_cell_text(
                table.cell(row_index, column),
                "☐",
                size=13,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
        set_cell_text(table.cell(row_index, 3), "________________________", size=9)
    comment = table.cell(7, 0).merge(table.cell(7, 3))
    set_cell_width(comment, TABLE_WIDTH)
    shade_cell(comment, PALE_BLUE)
    set_cell_text(
        comment,
        "Комментарий / обязательные формулировки: _______________________________________________________",
        size=9,
        color=MUTED,
    )


def build() -> Path:
    document = Document()
    configure_document(document)
    properties = document.core_properties
    properties.title = "SLIDE — открытые вопросы по стоимости и КП"
    properties.subject = "RFI по решениям для первой версии ценообразования"
    properties.author = "Raluma"
    properties.keywords = "SLIDE, RFI, стоимость, коммерческое предложение"

    add_masthead(document)
    add_approved_callout(document)

    questions = (
        (
            "Активировать ли наценку на отходы в первой версии?",
            (
                ("Не активировать; сохранить текущую настройку", ""),
                ("Активировать для всех ценовых позиций", "Процент: ____ %"),
                ("Активировать выборочно", "Категории / процент: __________"),
            ),
        ),
        (
            "Может ли дилер менять видимую скидку клиенту и в каких пределах?",
            (
                ("Нет; скидку задаёт система или менеджер", ""),
                ("Да; единый предел для всех дилеров", "Диапазон: 0–____ %"),
                ("Да; индивидуальный предел по дилеру или роли", "Кто задаёт: __________"),
            ),
        ),
        (
            "Кто вправе назначать разовую цену при отсутствии каталожной?",
            (
                ("Только superadmin / admin", ""),
                ("Сотрудник с правом управления ценами", "Роли: ______________"),
                ("Дилер в пределах отдельного лимита", "Лимит: ______________"),
            ),
        ),
        (
            "Кто вправе разрешать цену ниже минимальной маржи?",
            (
                ("Только superadmin", ""),
                ("Admin и superadmin", ""),
                ("Назначенный сотрудник с обязательным обоснованием", "Роли: ______________"),
            ),
        ),
        (
            "Хранить только последнюю редакцию КП или историю редакций?",
            (
                ("Только последнюю редакцию", ""),
                ("Только зафиксированные редакции", "Срок хранения: ________"),
                ("Полную историю черновиков и зафиксированных редакций", "Срок хранения: ________"),
            ),
        ),
        (
            "Относить ли покраску, изготовление и ручные операции к услугам?",
            (
                ("Да; все перечисленные позиции считать услугами", ""),
                ("Покраску и изготовление — услуги; ручные операции — отдельно", "Категория: ____________"),
                ("Нет; учитывать внутри готовой конструкции", ""),
            ),
        ),
    )
    for number, (question, options) in enumerate(questions, start=1):
        if number in {3, 6}:
            document.add_page_break()
        add_question(document, number, question, options)
    add_defaults_question(document)
    add_scope_question(document)

    closing = document.add_paragraph()
    set_paragraph(closing, before=10, after=2, line=1.0)
    run = closing.add_run("Решение согласовал(а): ")
    set_run_font(run, size=10, bold=True)
    run = closing.add_run("____________________________  Роль: __________________  Дата: ______________")
    set_run_font(run, size=10)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
