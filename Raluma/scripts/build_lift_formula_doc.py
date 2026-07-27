"""Build the customer-facing LIFT formula and document logic reference."""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(
    os.environ.get(
        "LIFT_FORMULA_DOC_OUTPUT",
        ROOT / "docs" / "ЛИФТ_формулы_и_логика_для_согласования.docx",
    )
)

ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
TEXT = "1F2933"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F3EC"
LIGHT_YELLOW = "FFF4CC"
LIGHT_RED = "FCE8E6"
WHITE = "FFFFFF"
GRID = "AAB4BE"

TABLE_WIDTH_DXA = 10062
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_cell_text(cell, text: str, *, bold=False, color=TEXT, size=9.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    font_size: float = 9.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            color=ACCENT_DARK,
            size=font_size,
        )
        set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], str(value), size=font_size)

    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_status_box(doc: Document, label: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), label, bold=True, color=ACCENT_DARK, size=9.5)
    set_cell_text(table.cell(0, 1), text, size=9.5)
    set_cell_shading(table.cell(0, 0), fill)
    set_cell_shading(table.cell(0, 1), fill)
    set_table_geometry(table, [1900, 8162])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_bullet(doc: Document, text: str, *, level=0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=TEXT)


def create_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    base_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = next(
        num
        for num in numbering.findall(qn("w:num"))
        if int(num.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(
        int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))
    ) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)

    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return new_num_id


def add_numbered(doc: Document, text: str, num_id: int | None = None) -> None:
    if num_id is None:
        previous = doc.paragraphs[-1] if doc.paragraphs else None
        previous_num_pr = (
            previous._p.pPr.numPr
            if previous is not None
            and previous._p.pPr is not None
            and previous._p.pPr.numPr is not None
            else None
        )
        if (
            previous is not None
            and previous.style.name == "List Number"
            and previous_num_pr is not None
        ):
            num_id = int(previous_num_pr.numId.val)
        else:
            num_id = create_numbering_instance(doc)

    paragraph = doc.add_paragraph(style="List Number")
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=TEXT)


def add_numbered_sequence(doc: Document, items: list[str]) -> None:
    num_id = create_numbering_instance(doc)
    for item in items:
        add_numbered(doc, item, num_id)


def add_formula(doc: Document, formula: str, note: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F8FB")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2 if note else 0)
    run = paragraph.add_run(formula)
    set_run_font(run, name="Consolas", size=9.5, color=ACCENT_DARK, bold=True)
    if note:
        note_paragraph = cell.add_paragraph()
        note_paragraph.paragraph_format.space_after = Pt(0)
        note_run = note_paragraph.add_run(note)
        set_run_font(note_run, size=9, color=MUTED, italic=True)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_keep_with_next(paragraph)
    run = paragraph.add_run(text)
    run.bold = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    # Named override to the preset baseline: A4 and 18 mm margins are used because
    # this is a Russian technical reference with dense formula tables.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, ACCENT, 18, 10),
        2: (13, ACCENT, 14, 7),
        3: (12, ACCENT_DARK, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_paragraph.add_run("RALUMA | ЛИФТ | расчётная логика")
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    add_page_number(section.footer.paragraphs[0])


def add_title_page(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("ЛИФТ")
    set_run_font(run, size=24, color=TEXT, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Формулы расчёта и логика документов")
    set_run_font(run, size=15, color=MUTED, bold=True)

    metadata = [
        ("Проект", "RALUMA"),
        ("Назначение", "Согласование реализованных расчётов и документов ЛИФТ"),
        ("Версия", "1.0 - реализация по Excel-источникам"),
        ("Дата", date.today().strftime("%d.%m.%Y")),
        ("Статус", "Реализовано и проверено. Спорные правила отмечены отдельно."),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=11, color=TEXT, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=11, color=TEXT)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run(
        "Документ является единым журналом реализованных формул ЛИФТ и используется "
        "для согласования расчётов с заказчиком. ЛИФТ реализован отдельным расчётным "
        "модулем и отдельным производственным листом. Формулы СЛАЙД не входят в "
        "область этой версии и при разработке ЛИФТ не изменялись."
    )
    set_run_font(run, size=11.5, color=TEXT)

    add_status_box(
        doc,
        "ПОДТВЕРЖДЕНО",
        "Правило однозначно извлечено из предоставленного Excel, реализовано и покрыто тестами.",
        LIGHT_GREEN,
    )
    add_status_box(
        doc,
        "СОГЛАСОВАТЬ",
        "В исходных Excel есть конфликт, готовое число вместо формулы или не определена граница условия.",
        LIGHT_YELLOW,
    )
    add_status_box(
        doc,
        "РЕАЛИЗОВАНО",
        "Расчёт используется в коде, проверен автоматическими тестами и ручным HTML-рендером документов.",
        LIGHT_BLUE,
    )

    add_heading(doc, "Принцип ведения документа", 2)
    add_bullet(doc, "Каждая формула хранится вместе с исходным файлом и назначением.")
    add_bullet(doc, "Противоречия не исправляются догадкой: они выносятся на согласование.")
    add_bullet(doc, "Реализованные формулы закреплены тест-кейсами и журналом проверки.")
    add_bullet(doc, "Расчёт ЛИФТ реализуется отдельным модулем и не использует slide_calc.py.")


def add_sources(doc: Document) -> None:
    add_heading(doc, "1. Исходные материалы", 1)
    rows = [
        ["1", "Lift 2ух стекло (2).xlsx", "2 панели, стекло 8 мм", "Извлечено"],
        ["2", "Lift 2ух стеклопакет (2).xlsx", "2 панели, стеклопакет 20 мм", "Извлечено"],
        ["3", "Lift 3ех стекло (2).xlsx", "3 панели, стекло 8 мм", "Извлечено"],
        ["4", "Lift 3ех стеклопакет (2).xlsx", "3 панели, стеклопакет 20 мм", "Извлечено"],
        ["5", "Lift 4ех- стекло.xlsx", "4 панели, обычное открывание", "Есть конфликт 1 мм"],
        [
            "6",
            "Lift_4ех_стекло_глух_вверху_и_внизу (2).xlsx",
            "4 панели, верх/низ глухие",
            "Извлечено",
        ],
        ["7", "Lift 4ех- стеклопакет.xlsx", "4 панели, стеклопакет", "Есть конфликт 1 мм"],
        [
            "8",
            "Lift_4ех_стеклопакет_глух_вверху_и_внизу.xlsx",
            "4 панели, стеклопакет, верх/низ глухие",
            "Извлечено",
        ],
        ["9", "фурнитура ЛИФТ.xlsx", "Фурнитура, цепь, привод, крутящий момент", "Извлечено"],
        ["10", "картинки ЛИФТ-...zip", "39 PNG по артикулам RL/RU", "Проверено визуально"],
    ]
    add_table(
        doc,
        ["№", "Источник", "Назначение", "Статус"],
        rows,
        [500, 4000, 3650, 1912],
        font_size=8.7,
    )
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "Примечание: Excel-файлы содержат около 863 формул и более 470 встроенных "
        "изображений. Часть встроенных WMF не читается библиотекой, поэтому рабочие "
        "изображения берутся из отдельного ZIP."
    )
    set_run_font(run, size=10, color=MUTED, italic=True)


def add_notation_and_architecture(doc: Document) -> None:
    add_heading(doc, "2. Переменные и архитектурные ограничения", 1)
    add_table(
        doc,
        ["Обозначение", "Смысл", "Единица"],
        [
            ["W", "Габаритная ширина секции", "мм"],
            ["H", "Габаритная высота секции", "мм"],
            ["P", "Количество панелей: 2, 3 или 4", "шт"],
            ["Q", "Количество одинаковых секций", "шт"],
            ["F", "Количество глухих панелей", "шт"],
            ["M", "Количество подвижных панелей: P - F", "шт"],
            ["T", "Крутящий момент привода", "Н·м"],
        ],
        [1600, 6762, 1700],
    )

    add_heading(doc, "Обязательная изоляция от СЛАЙД", 2)
    add_numbered(doc, "Создать отдельный backend-модуль engine/lift_calc.py.")
    add_numbered(doc, "Не менять engine/slide_calc.py и шаблон section_sheet.html ради ЛИФТ.")
    add_numbered(
        doc,
        "Расчёт строить от физических панелей, а документы группировать только после расчёта.",
    )
    add_numbered(
        doc,
        "Для ЛИФТ использовать отдельный шаблон производственного листа lift_section_sheet.html.",
    )
    add_numbered(
        doc,
        "В общепроектных документах выбирать калькулятор по system, не смешивая модели результатов.",
    )

    doc.add_page_break()
    add_heading(doc, "Предлагаемая расчётная модель", 2)
    add_table(
        doc,
        ["Сущность", "Минимальные поля"],
        [
            [
                "LiftPhysicalPanel",
                "index, fixed/moving, filling_type, filling_width, filling_height, glued_width, glued_height",
            ],
            [
                "LiftProfileItem",
                "article, name, length_mm, qty, image, painted, note",
            ],
            [
                "LiftHardwareItem",
                "article, name, qty, unit, image, note",
            ],
            [
                "LiftCalcResult",
                "panels, profiles, hardware, torque, drives, warnings, diagram namespaces",
            ],
        ],
        [2550, 7512],
        font_size=9,
    )


def add_common_profiles(doc: Document) -> None:
    add_heading(doc, "3. Общие профили рамы", 1)
    add_status_box(
        doc,
        "ПОДТВЕРЖДЕНО",
        "Следующие формулы повторяются во всех восьми расчётных вариантах.",
        LIGHT_GREEN,
    )
    rows = [
        ["RL101-1", "W - 6", "3Q", "Общий профиль рамы"],
        ["RL101", "W - 6", "Q", "Общий профиль рамы"],
        ["RL102", "H - 161", "Q", "Вертикальный профиль"],
        ["RL103", "H - 161", "Q", "Вертикальный профиль"],
        ["RL103-1", "H - 161", "2Q", "Вертикальный профиль"],
        ["RL103-2", "H - 161", "Q", "Вертикальный профиль"],
        ["RL104", "W - 155", "Q", "Внутреннее/верхнее положение"],
        ["RL104", "W - 62", "Q", "Второе положение: (W - 6) - 56"],
    ]
    add_table(
        doc,
        ["Артикул", "Длина, мм", "Количество", "Примечание"],
        rows,
        [1700, 2500, 1700, 4162],
    )

    add_heading(doc, "Правило объединения нарезки", 2)
    add_formula(
        doc,
        "Ключ объединения = (артикул, округлённая длина, цвет/покрытие, операция)",
        "Одинаковый артикул одинаковой длины выводится одной строкой с суммарным количеством.",
    )
    add_bullet(
        doc,
        "Позиции с разной длиной не объединяются, даже если артикул совпадает.",
    )
    add_bullet(
        doc,
        "Повтор одной и той же позиции в разных местах Excel не должен создавать дубли в ПЛ.",
    )


def add_panel_formulas(doc: Document) -> None:
    add_heading(doc, "4. Заполнение и панели при склейке", 1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "Сначала рассчитывается каждая физическая панель. Таблицы стекла/стеклопакета и "
        "панелей при склейке формируются из одного массива, чтобы схема, ПЛ и заявка "
        "на стекло не расходились."
    )
    set_run_font(run, size=11, color=TEXT)

    doc.add_page_break()
    add_heading(doc, "4.1. Две панели - стекло 8 мм", 2)
    add_formula(doc, "Ширина подвижного заполнения = W - 133")
    add_formula(doc, "Ширина глухого заполнения = W - 135")
    add_formula(
        doc,
        "База по высоте = (H - 213 - 11,5) / 2 - 1",
        "Высота подвижного заполнения = База + 11,5; высота глухого = База.",
    )
    add_formula(
        doc,
        "Панель при склейке = (ширина заполнения + 47) x (высота заполнения + 46)",
    )

    add_heading(doc, "4.2. Две панели - стеклопакет 20 мм", 2)
    add_bullet(
        doc,
        "Геометрия заполнений и панелей при склейке совпадает с двухпанельным стеклом.",
    )
    add_bullet(
        doc,
        "Подвижная панель: стеклопакет 6зак-8AR-6зак; глухая панель: Пеноплекс 20 мм.",
    )
    add_bullet(
        doc,
        "Профили заполнения используют линейку RL123/RL122/RL1241/RL1211.",
    )

    add_heading(doc, "4.3. Три панели - стекло 8 мм", 2)
    add_formula(doc, "Ширина двух панелей = W - 135")
    add_formula(doc, "Ширина подвижной панели = W - 133")
    add_formula(
        doc,
        "База по высоте = (H - 216 - 11,5) / 3 - 1",
        "Одна подвижная панель получает База + 11,5; остальные получают База.",
    )
    add_formula(
        doc,
        "Панель при склейке = (ширина заполнения + 47) x (высота заполнения + 46)",
    )

    add_heading(doc, "4.4. Три панели - стеклопакет 20 мм", 2)
    add_formula(doc, "Ширина всех заполнений = W - 133")
    add_formula(
        doc,
        "База по высоте = (H - 216 - 11,5) / 3 - 1",
        "Одна панель получает База + 11,5; остальные получают База.",
    )
    add_formula(
        doc,
        "Панель при склейке = (ширина заполнения + 47) x (высота заполнения + 46)",
    )

    add_heading(doc, "4.5. Четыре панели - обычное открывание", 2)
    add_status_box(
        doc,
        "СОГЛАСОВАТЬ",
        "Исходные Excel расходятся на 1 мм: часть листов использует W - 134 без финального -1, часть W - 135.",
        LIGHT_YELLOW,
    )
    add_formula(
        doc,
        "Вариант A: ширина заполнения = W - 134",
        "Использован в первых листах Lift 4ех- стекло.xlsx.",
    )
    add_formula(
        doc,
        "Вариант B: ширина заполнения = W - 135",
        "Использован в третьем листе и в файле стеклопакета.",
    )
    add_formula(
        doc,
        "База по высоте = (H - 160 - 110 - 11) / 4 [в части листов дополнительно -1]",
    )
    add_formula(
        doc,
        "Панель при склейке = (ширина заполнения + 48) x (высота заполнения + 46)",
    )

    add_heading(doc, "4.6. Четыре панели - верх и низ глухие", 2)
    add_formula(doc, "Ширина заполнения = W - 133")
    add_formula(doc, "База по высоте = (H - 161 - 116 - 11) / 4 - 1")
    add_formula(
        doc,
        "Панель при склейке = (ширина заполнения + 47) x (высота заполнения + 47)",
    )
    add_bullet(
        doc,
        "Дополнительный RL103-2: длина = длина RL113 + 70 мм, количество 1; операция - срезать под уплотнение 4-й створки.",
    )

    add_heading(doc, "4.7. Профили панелей", 2)
    add_table(
        doc,
        ["Вариант", "Основные артикулы", "Правило"],
        [
            [
                "Стекло 8 мм",
                "RL113, RL112, RL115, RL114, RL105",
                "Длины выводятся от рассчитанных физических заполнений.",
            ],
            [
                "Стеклопакет 20 мм",
                "RL123, RL122, RL1241, RL1211, RL105",
                "Сохранять разные смещения -45/-46 из Excel.",
            ],
            [
                "3 панели",
                "RL113/RL123 могут повторяться",
                "Одинаковые длины объединять после расчёта.",
            ],
            [
                "4 панели",
                "RL105",
                "В Excel встречаются готовые 1903/2823; общая формула не найдена.",
            ],
        ],
        [2200, 3200, 4662],
        font_size=9,
    )


def add_hardware(doc: Document) -> None:
    add_heading(doc, "5. Фурнитура ЛИФТ", 1)
    add_status_box(
        doc,
        "ПОДТВЕРЖДЕНО",
        "Базовые количества извлечены из файла «фурнитура ЛИФТ.xlsx».",
        LIGHT_GREEN,
    )
    rows = [
        ["RL201", "Угол соединения рамы", "2Q"],
        ["RL001", "Угол панели для стекла 8 мм", "4 x P x Q"],
        ["RL011", "Угол панели для стеклопакета 20 мм", "4 x P x Q"],
        ["RL203", "Заглушка вала", "(2 - число приводов) x Q"],
        ["RL206", "Звёздочка цепи", "2Q"],
        ["RL207", "Подшипник", "(2 - число приводов) x Q"],
        ["RL2096", "Адаптер привода", "число приводов x Q"],
        ["RL2097", "Крепление привода", "число приводов x Q"],
        ["RL2087", "Пульт 1 канал", "общее количество на проект"],
        ["RL2088", "Пульт 6 каналов", "общее количество на проект"],
        ["RL2092", "Кнопка", "Q для секций с кнопкой"],
        ["RL005", "Успокоитель цепи", "2Q"],
        ["RL002", "Заглушка панели", "2Q"],
        ["RU1039", "Наклейка", "Q"],
        ["RL150", "Инструкция ЛИФТ", "Требует подтверждения: 1 на проект или Q"],
    ]
    add_table(
        doc,
        ["Артикул", "Наименование", "Количество"],
        rows,
        [1700, 5262, 3100],
        font_size=9,
    )

    add_heading(doc, "Боковые крышки по вводу кабеля", 2)
    add_bullet(doc, "RL20901 - крышка с подшипником слева.")
    add_bullet(doc, "RL20902 - крышка под мотор справа.")
    add_bullet(doc, "RL20903 - крышка с подшипником справа.")
    add_bullet(doc, "RL20904 - крышка под мотор слева.")
    add_bullet(
        doc,
        "Выбор пары зависит от стороны ввода кабеля и количества приводов.",
    )

    add_heading(doc, "Цепь RL210", 2)
    add_formula(
        doc,
        "2 панели: L = округление вверх до 100 мм [H/2 + 350 + pi x 110 / 2 + 200]",
    )
    add_formula(
        doc,
        "3 панели: L = округление вверх до 100 мм [2H/3 + 350 + pi x 110 / 2 + 200]",
    )
    add_formula(
        doc,
        "4 панели: L = округление вверх до 100 мм [3H/4 + 350 + pi x 110 / 2 + 200]",
        "Количество цепей: 2Q.",
    )

    add_heading(doc, "Уплотнители и крепёж", 2)
    add_table(
        doc,
        ["Позиция", "Формула/количество", "Статус"],
        [
            [
                "RU004, щётка 7x6",
                "ROUNDUP((D1 x (P x 2) x Q) / 1000, 0)",
                "Нужно определить смысл D1",
            ],
            ["RU006, щётка 7x12", "ROUNDUP((H x 12 x Q) / 1000, 0)", "Извлечено"],
            ["DIN7982 4,2x16", "6Q", "Извлечено"],
            ["DIN7982 4,2x80", "2Q", "Извлечено"],
            ["DIN7982 4,8x16", "6Q", "По комментарию Excel"],
            ["DIN7504O 4,8x16", "6Q", "По комментарию Excel"],
            ["DIN7985 M4x20", "8Q", "Извлечено"],
            ["DIN125, шайба d4", "8Q", "Извлечено"],
            ["DIN985 M4", "8Q", "Извлечено"],
            ["DIN965 M6x10", "4Q при P >= 3", "Извлечено"],
            ["DIN965 M6x20", "8Q при P < 3; иначе 4 x (P - 1) x Q", "Извлечено"],
            [
                "DIN7504O 3,9x13",
                "H / 300 x 3 x Q + 5Q",
                "Нужно подтвердить округление",
            ],
        ],
        [2700, 4700, 2662],
        font_size=8.7,
    )


def add_torque(doc: Document) -> None:
    add_heading(doc, "6. Крутящий момент и привод", 1)
    add_formula(
        doc,
        "M = P - F",
        "Количество подвижных панелей равно общему числу панелей минус глухие.",
    )
    add_formula(
        doc,
        "Вес = площадь наибольшей панели x M x эквивалентная толщина x 2,5 x 1,1",
        "1,1 - запас 10%; для стекла 8 мм толщина 8, для стеклопакета 20 мм в Excel используется 12.",
    )
    add_formula(
        doc,
        "T = Вес x 9,81 x 51 / 1000",
        "Результат выводится с точностью 0,1 Н·м.",
    )
    add_table(
        doc,
        ["Условие", "Пульт ДУ", "Кнопка"],
        [
            [
                "T < 80 Н·м",
                "1 x RL2085 (радиосвязь) на секцию",
                "1 x RL2095 (фазный) на секцию",
            ],
            [
                "T > 80 Н·м",
                "2 x RL2095 + 1 x RL2098 на секцию",
                "2 x RL2095 + 1 x RL2098 на секцию",
            ],
            [
                "T > 160 Н·м",
                "Красное предупреждение о превышении допустимого усилия",
                "Красное предупреждение о превышении допустимого усилия",
            ],
        ],
        [2200, 3931, 3931],
        font_size=9,
    )
    add_status_box(
        doc,
        "СОГЛАСОВАТЬ",
        "Excel задаёт только условия <80, >80 и >160. Нужно определить поведение ровно при 80 и 160 Н·м.",
        LIGHT_YELLOW,
    )


def add_documents_plan(doc: Document) -> None:
    add_heading(doc, "7. Реализованные документы ЛИФТ", 1)
    add_status_box(
        doc,
        "РЕАЛИЗОВАНО",
        "Excel используется только как источник формул, артикулов и зависимостей. "
        "Его перегруженные листы не копируются в интерфейс или PDF. Документы "
        "формируются отдельными аккуратными HTML/PDF-шаблонами в стиле RALUMA. "
        "Этот Word фиксирует реализованную расчётную логику и вопросы согласования.",
        LIGHT_BLUE,
    )
    add_heading(doc, "7.1. Производственный лист", 2)
    add_numbered(
        doc,
        "Отдельный файл на каждую секцию ЛИФТ; используется lift_section_sheet.html, шаблон СЛАЙД не изменён.",
    )
    add_numbered(
        doc,
        "Страница 1: параметры секции, вид из помещения, вертикальная кинематика, схема склейки и таблица физических панелей.",
    )
    add_numbered(
        doc,
        "Страница 2: нарезка профилей с объединением только одинаковых артикулов и длин.",
    )
    add_numbered(
        doc,
        "Страница 3: фурнитура, крепёж, привод, вес, крутящий момент, предупреждения и комментарии.",
    )
    add_numbered(
        doc,
        "Все три схемы строятся из того же расчёта физических панелей, что и таблицы.",
    )

    add_heading(doc, "7.2. Заявка на стекло/заполнение", 2)
    add_bullet(
        doc,
        "Подключена к общепроектной заявке; размеры берутся из физических панелей ЛИФТ.",
    )
    add_bullet(
        doc,
        "Строить из физических панелей, а не из сгруппированных строк ПЛ.",
    )
    add_bullet(
        doc,
        "Группировать только одинаковые тип, ширину, высоту и примечание.",
    )
    add_bullet(
        doc,
        "Для стеклопакета сохранять полное обозначение 6зак-8AR-6зак.",
    )

    add_heading(doc, "7.3. Заявка на покраску", 2)
    add_bullet(
        doc,
        "Подключена к общепроектной заявке; каждый цвет выводится на своей странице.",
    )
    add_bullet(
        doc,
        "Брать только окрашиваемые профили ЛИФТ из рассчитанной нарезки.",
    )
    add_bullet(
        doc,
        "Группировать по цвету, артикулу, чистовому размеру и операции.",
    )
    add_bullet(
        doc,
        "Использовать изображения ЛИФТ из ZIP, привязанные к артикулам каталога.",
    )
    add_bullet(
        doc,
        "Ручные строки проекта добавлять в конец соответствующего цвета и включать в общий итог.",
    )

    add_heading(doc, "7.4. Накладная", 2)
    add_bullet(
        doc,
        "ЛИФТ подключён к общепроектной накладной RALUMA.",
    )
    add_bullet(
        doc,
        "Заполнение ЛИФТ брать из lift_filling_type/lift_filling_custom, а не из glass_type.",
    )
    add_bullet(
        doc,
        "Не объединять конструкции с разным числом панелей, открыванием или управлением.",
    )
    add_bullet(
        doc,
        "Количество мест рассчитывать отдельно по группам профилей, заполнений и фурнитуры.",
    )

    add_heading(doc, "7.5. Интеграция API", 2)
    add_bullet(
        doc,
        "Section preview/calc/pdf выбирает calculate_slide или calculate_lift по system.",
    )
    add_bullet(
        doc,
        "Предпросмотр и PDF производственного листа ЛИФТ используют один контекст расчёта.",
    )
    add_bullet(
        doc,
        "Гостевой и авторизованный режимы используют один расчёт и один контекст.",
    )


def add_diagrams_and_assets(doc: Document) -> None:
    add_heading(doc, "8. Схемы и изображения", 1)
    add_table(
        doc,
        ["Схема", "Источник данных", "Требование"],
        [
            [
                "Вид из помещения",
                "LiftPhysicalPanel[]",
                "Пропорционально W/H; 7 вариантов, направление и глухие панели из параметров секции.",
            ],
            [
                "Вертикальная кинематика",
                "Панели + RL105 + точки подвеса",
                "Показать последовательность панелей и связи, не копировать растр из Excel.",
            ],
            [
                "Сечение",
                "Профили + filling/glued panel",
                "Показывать реальные рассчитанные размеры заполнения и панели при склейке.",
            ],
        ],
        [2500, 2900, 4662],
        font_size=9,
    )

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "ZIP содержит 39 PNG: RL001, RL002, RL005, RL011, RL101/101-1, "
        "RL102, RL103/103-1/103-2, RL104, RL105, RL112-RL1241, RL201-RL210, "
        "RL2085/RL2095, RL2096-RL2098, гайку и шайбу. Изображения перенесены "
        "в assets, привязаны к артикулам каталога и используются в документах ЛИФТ."
    )
    set_run_font(run, size=10.5, color=TEXT)


def add_qa_plan(doc: Document) -> None:
    add_heading(doc, "9. Реализация и проверка", 1)
    steps = [
        "Созданы фикстуры и расчётные тесты для всех восьми Excel-вариантов.",
        "Реализован отдельный engine/lift_calc.py: панели, склейка, профили, фурнитура, цепь, привод и крутящий момент.",
        "Позиции одинакового артикула объединяются только при одинаковой длине и покрытии.",
        "Изображения ЛИФТ перенесены в assets и подключены к каталогу.",
        "Создан отдельный трёхстраничный производственный лист ЛИФТ с тремя схемами.",
        "ЛИФТ подключён к заявке на стекло, заявке на покраску и накладной.",
        "Проверены гостевой и авторизованный режимы, копирование и шаблоны секций.",
        "Выполнен ручной визуальный контроль HTML всех четырёх документов.",
    ]
    for step in steps:
        add_numbered(doc, step)

    add_heading(doc, "Минимальная матрица тестов", 2)
    add_table(
        doc,
        ["Панели", "Заполнение", "Открывание", "Документы"],
        [
            ["2", "Стекло 8 мм", "Вверх и вниз", "calc, ПЛ, стекло, покраска, накладная"],
            ["2", "Стеклопакет 20 мм", "Вверх и вниз", "calc, ПЛ, стекло, покраска, накладная"],
            ["3", "Стекло 8 мм", "Вверх и вниз", "calc, ПЛ, стекло, покраска, накладная"],
            ["3", "Стеклопакет 20 мм", "Вверх и вниз", "calc, ПЛ, стекло, покраска, накладная"],
            ["4", "Стекло 8 мм", "Вверх, вниз, верх/низ глухие", "все документы"],
            ["4", "Стеклопакет 20 мм", "Вверх, вниз, верх/низ глухие", "все документы"],
        ],
        [1200, 2500, 2762, 3600],
        font_size=8.8,
    )
    add_status_box(
        doc,
        "РЕЗУЛЬТАТ ПРОВЕРОК",
        "Backend: 363 passed, 2 skipped. Ruff, TypeScript typecheck, frontend build, "
        "smoke:diagrams и smoke:lift-input пройдены. Визуальный HTML-контроль ПЛ, "
        "заявки на стекло, покраску и накладной пройден.",
        LIGHT_GREEN,
    )
    add_status_box(
        doc,
        "ОСТАЛОСЬ ПЕРЕД РЕЛИЗОМ",
        "Проверить финальные PDF внутри Docker/VPS с установленным WeasyPrint и "
        "получить подтверждение заказчика по правилам раздела 10.",
        LIGHT_YELLOW,
    )


def add_open_questions(doc: Document) -> None:
    add_heading(doc, "10. Реализованные допущения для согласования", 1)
    questions = [
        (
            "Q1",
            "Четыре панели: ширина заполнения",
            "Сейчас: стекло 8 мм = W - 134, стеклопакет 20 мм = W - 135. Подтвердить.",
        ),
        (
            "Q2",
            "Четыре панели: высота",
            "Реализованы отдельные формулы обычного и варианта с глухими верхом/низом. Подтвердить финальные вычеты 1 мм.",
        ),
        (
            "Q3",
            "RL105 для четырёх панелей",
            "Пока используются готовые длины 1903 и 2823 из Excel. Нужна общая формула.",
        ),
        (
            "Q4",
            "Граница привода",
            "Сейчас T <= 80 Н·м: 1 привод; T > 80: 2 привода; T > 160: предупреждение. Подтвердить.",
        ),
        (
            "Q5",
            "RU004 7x6",
            "Сейчас количество считается от ширины, числа панелей и секций. Подтвердить вместо неоднозначного D1 из Excel.",
        ),
        (
            "Q6",
            "DIN7504O 3,9x13",
            "Сейчас H / 300 округляется вверх до целого перед расчётом количества. Подтвердить.",
        ),
        (
            "Q7",
            "Инструкция RL150",
            "Сейчас добавляется одна инструкция на расчёт секции. Подтвердить: на проект или на секцию.",
        ),
        (
            "Q8",
            "Стеклопакет",
            "Сейчас эквивалентная толщина для веса и момента = 12 мм. Подтвердить.",
        ),
        (
            "Q9",
            "PDF",
            "Перед релизом проверить все документы в Docker/VPS с WeasyPrint.",
        ),
        (
            "Q10",
            "Чертежи",
            "Если в проект добавляются отдельные PDF-чертежи, подтвердить порядок их присоединения к документам.",
        ),
    ]
    add_table(
        doc,
        ["ID", "Тема", "Что подтвердить"],
        [list(row) for row in questions],
        [800, 2900, 6362],
        font_size=9,
    )
    add_status_box(
        doc,
        "ВАЖНО",
        "Расчёт реализован по указанным допущениям и покрыт тестами, но Q1-Q8 "
        "нужно подтвердить у заказчика до окончательной фиксации производственных правил.",
        LIGHT_RED,
    )


def add_change_log(doc: Document) -> None:
    add_heading(doc, "11. Журнал изменений", 1)
    add_table(
        doc,
        ["Версия", "Дата", "Изменение", "Статус"],
        [
            [
                "0.1",
                "27.07.2026",
                "Систематизированы 9 Excel, формулы заполнений, профилей, фурнитуры, цепи и привода; составлен план документов.",
                "На согласование",
            ],
            [
                "1.0",
                date.today().strftime("%d.%m.%Y"),
                "Реализованы отдельный расчёт ЛИФТ, каталог, трёхстраничный ПЛ, заявка на стекло, покраску, накладная и тесты.",
                "Реализовано",
            ],
        ],
        [1100, 1600, 5762, 1600],
        font_size=9,
    )
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "Следующая версия зафиксирует ответы заказчика на Q1-Q8 и результат "
        "финальной PDF-проверки внутри Docker/VPS."
    )
    set_run_font(run, size=10, color=MUTED, italic=True)


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_sources(doc)
    add_notation_and_architecture(doc)
    add_common_profiles(doc)
    add_panel_formulas(doc)
    add_hardware(doc)
    add_torque(doc)
    add_documents_plan(doc)
    add_diagrams_and_assets(doc)
    add_qa_plan(doc)
    add_open_questions(doc)
    add_change_log(doc)
    doc.core_properties.title = "ЛИФТ - формулы расчёта и логика документов"
    doc.core_properties.subject = "RALUMA: согласование расчётов ЛИФТ"
    doc.core_properties.author = "RALUMA"
    doc.core_properties.keywords = "ЛИФТ, формулы, производственный лист, RALUMA"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
