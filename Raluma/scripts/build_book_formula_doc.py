"""Build the unified BOOK formula register for calculator approval.

Design system: ``compact_reference_guide`` with the ``memo_masthead`` opening
pattern. The document is intentionally generated from the approved source
hierarchy: specification -> new Excel workbooks -> legacy application.
"""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "docs"
    / "КНИЖКА_формулы_и_открытые_вопросы_для_согласования.docx"
)
DOCUMENT_DATE = "30.07.2026"

# compact_reference_guide, exact preset tokens.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT = "Calibri"
MONO = "Consolas"
INK = "1F2933"
MUTED = "5F6B76"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRID = "AAB4BE"
HEADER_FILL = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E8F3EC"
LIGHT_YELLOW = "FFF4CC"
LIGHT_RED = "FCE8E6"
FORMULA_FILL = "F5F8FB"
WHITE = "FFFFFF"

STATUS_FILL = {
    "ПОДТВЕРЖДЕНО": LIGHT_GREEN,
    "ПРЕДВАРИТЕЛЬНО": LIGHT_YELLOW,
    "ЗАБЛОКИРОВАНО": LIGHT_RED,
    "ОТКРЫТЫЙ ВОПРОС": LIGHT_GRAY,
}


def set_run_font(
    run,
    *,
    name: str = FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in CELL_MARGINS_DXA.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = GRID) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table width mismatch: {sum(widths_dxa)} != {TABLE_WIDTH_DXA}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_text(
    cell,
    text: str,
    *,
    size: float = 8.5,
    bold: bool = False,
    color: str = INK,
    font: str = FONT,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(str(text))
    set_run_font(run, name=font, size=size, color=color, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    font_size: float = 8.4,
    status_col: int | None = None,
    mono_cols: set[int] | None = None,
    center_cols: set[int] | None = None,
) -> None:
    mono_cols = mono_cols or set()
    center_cols = center_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            size=font_size,
            bold=True,
            color=DARK_BLUE,
            align=(
                WD_ALIGN_PARAGRAPH.CENTER
                if index in center_cols
                else WD_ALIGN_PARAGRAPH.LEFT
            ),
        )
        set_cell_shading(table.rows[0].cells[index], HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell_text(
                row.cells[index],
                value,
                size=font_size,
                font=MONO if index in mono_cols else FONT,
                align=(
                    WD_ALIGN_PARAGRAPH.CENTER
                    if index in center_cols
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
            )
            if status_col == index:
                status = str(value).split("\n")[-1].strip().upper()
                for label, fill in STATUS_FILL.items():
                    if label in status:
                        set_cell_shading(row.cells[index], fill)
                        break
        prevent_row_split(row)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold_prefix: str | None = None,
    style: str | None = None,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)


def add_formula(doc: Document, formula: str) -> None:
    paragraph = doc.add_paragraph(style="Formula")
    run = paragraph.add_run(formula)
    set_run_font(run, name=MONO, size=10, color=DARK_BLUE, bold=True)


def add_callout(
    doc: Document,
    title: str,
    body: str,
    fill: str,
    *,
    trailing_space: bool = True,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    title_run = paragraph.add_run(f"{title}\n")
    set_run_font(title_run, size=10.5, bold=True, color=DARK_BLUE)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, size=9.5, color=INK)
    set_cell_shading(cell, fill)
    prevent_row_split(table.rows[0])
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    if trailing_space:
        doc.add_paragraph()


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    formula_style = doc.styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    formula_style.font.name = MONO
    formula_style.font.size = Pt(10)
    formula_style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    formula_style.paragraph_format.left_indent = Inches(0.18)
    formula_style.paragraph_format.right_indent = Inches(0.18)
    formula_style.paragraph_format.space_before = Pt(4)
    formula_style.paragraph_format.space_after = Pt(8)
    formula_style.paragraph_format.keep_together = True
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), FORMULA_FILL)
    formula_style.element.get_or_add_pPr().append(shading)

    citation_style = doc.styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    citation_style.font.name = FONT
    citation_style.font.size = Pt(8.5)
    citation_style.font.italic = True
    citation_style.font.color.rgb = RGBColor.from_string(MUTED)
    citation_style.paragraph_format.space_before = Pt(4)
    citation_style.paragraph_format.space_after = Pt(4)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run("КНИЖКА · ЕДИНЫЙ РЕЕСТР ФОРМУЛ")
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_run = footer_paragraph.add_run("Страница ")
    set_run_font(footer_run, size=8.5, color=MUTED)
    add_page_field(footer_paragraph)


def add_opening(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    kicker_run = kicker.add_run("ТЕХНИЧЕСКИЙ РЕЕСТР · ЭТАП 1")
    set_run_font(kicker_run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("КНИЖКА")
    set_run_font(title_run, size=25, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run(
        "Формулы, источники, статусы и открытые вопросы для согласования калькулятора"
    )
    set_run_font(subtitle_run, size=13, color=MUTED)

    metadata = [
        ("Назначение:", "расчёт прямых секций, физических панелей, профилей и фурнитуры"),
        ("Приоритет:", "ТЗ КНИЖКА → новые Excel → восстановленные данные старой программы"),
        ("Точность:", "0,1 мм; производственное округление до целых не применяется"),
        ("Версия:", f"{DOCUMENT_DATE} · первый этап"),
        ("Документы:", "ПЛ, стекло, покраска и накладная — следующий пакет"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label} ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)
    doc.add_paragraph()

    add_callout(
        doc,
        "Главное правило источников",
        "Если ТЗ содержит формулу для позиции или конфигурации, формула Excel и "
        "старой программы для той же области не применяется. Excel используется "
        "только при отсутствии формулы в ТЗ; legacy — только при отсутствии обоих.",
        LIGHT_GREEN,
    )


def source_rows() -> list[list[str]]:
    return [
        [
            "1",
            "ТЗ КНИЖКА.docx",
            "Геометрия прямой секции, параметры формы, схемы, компенсаторы, "
            "отдельные правила углов/отверстий",
            "Высший приоритет",
        ],
        [
            "2",
            "Расчет_книжки_29_07_26_прямые_секции.xlsx",
            "Высотные вычеты, профильные позиции, сборочные зависимости",
            "Только если формулы нет в ТЗ",
        ],
        [
            "2",
            "Список фурнитуры книжка.xlsx",
            "Текстовые формулы количества и этап отгрузки каждой позиции",
            "Только если формулы нет в ТЗ",
        ],
        [
            "3",
            "Старая программа AppGlass / восстановленный контур",
            "Совместимость старых полей и резервные сведения",
            "Только если нет ТЗ и новых Excel",
        ],
    ]


def variable_rows() -> list[list[str]]:
    return [
        ["W", "Ширина секции", "мм", "Ввод пользователя"],
        ["H", "Высота секции", "мм", "Ввод пользователя"],
        ["P", "Количество базовых складных панелей", "шт.", "2…6"],
        ["Pфиз", "Количество физических панелей", "шт.", "P + доп. глухая панель"],
        ["q", "Количество одинаковых секций", "шт.", "> 0"],
        ["G", "Ширина стекла стандартной панели", "мм", "Точность 0,1"],
        ["Gp", "Ширина стекольного профиля", "мм", "G + 3; у угловой двери предв. G + 23"],
        ["D", "Количество крайних дверей", "шт.", "0, 1 или 2"],
        ["M", "Количество обычных подвижных панелей", "шт.", "Без крайних и дополнительных дверей"],
        ["A", "Количество угловых стыков", "шт.", "0…2, предварительно"],
        ["Jпрям", "Стыки прямой направляющей > 6000 мм", "шт.", "ceil(W / 6000) − 1"],
        ["C", "Количество компенсаторов", "шт.", "1 нижний/верхний; 2 — оба"],
    ]


def geometry_rows() -> list[list[str]]:
    return [
        [
            "Ширина стекла",
            "(W − 11,5 − 11,5 − 3 × (P − 1)) / P",
            "Прямая секция без углов и дополнительных панелей",
            "ТЗ, стр. 9",
            "tz\nПОДТВЕРЖДЕНО",
        ],
        [
            "Стекольный профиль",
            "Gp = G + 3",
            "Каждая физическая панель прямой секции",
            "ТЗ, стр. 9",
            "tz\nПОДТВЕРЖДЕНО",
        ],
        [
            "Высота стекла",
            "Hg = H − 135",
            "Прямые листы Excel; вычет требует согласования",
            "Excel C24/C25, C25/C26",
            "excel\nПРЕДВАРИТЕЛЬНО",
        ],
        [
            "Высота панели при склейке",
            "Hp = Hg + 33",
            "Прямые листы Excel",
            "Excel C27/C28, C29/C30",
            "excel\nПРЕДВАРИТЕЛЬНО",
        ],
        [
            "Край у угла",
            "Lкрай или Rкрай = 27 вместо 11,5",
            "Угловая конструкция",
            "ТЗ, стр. 9",
            "tz\nПРЕДВАРИТЕЛЬНО",
        ],
        [
            "Профиль угловой двери",
            "Gp = G + 23",
            "Дверь рядом с углом",
            "ТЗ, стр. 9",
            "tz\nПРЕДВАРИТЕЛЬНО",
        ],
        [
            "Доп. глухая панель",
            "Gобыч = (доступный пролёт − ΣGзаданных) / Pобыч",
            "Ширина доп. панели задана пользователем",
            "Расширение ТЗ для UI",
            "tz\nПРЕДВАРИТЕЛЬНО",
        ],
        [
            "Доп. двигающаяся дверь",
            "Gдоп = Wдоп − 3; 0 < Wдоп ≤ 850",
            "Дверь заменяет обычную подвижную панель",
            "ТЗ + пустой блок Excel",
            "tz/excel\nПРЕДВАРИТЕЛЬНО",
        ],
        [
            "Контроль размеров",
            "Все рассчитанные длины и размеры > 0",
            "Все конфигурации",
            "Правило калькулятора этапа 1",
            "tz\nПОДТВЕРЖДЕНО",
        ],
    ]


def profile_rows() -> list[list[str]]:
    return [
        [
            "RBP001",
            "Направляющий профиль",
            "L = W; qty = 2 × q",
            "Верх и низ секции",
            "Excel\nПОДТВЕРЖДЕНО",
        ],
        [
            "RBP003",
            "Компенсирующий профиль",
            "L = W; qty = C × q",
            "Нижний, верхний или оба",
            "ТЗ + Excel\nПОДТВЕРЖДЕНО",
        ],
        [
            "RBP002",
            "Стекольный профиль",
            "L = G + 3; qty = 2 × q на панель",
            "Прямая физическая панель",
            "ТЗ + Excel\nПОДТВЕРЖДЕНО",
        ],
        [
            "RBP002",
            "Стекольный профиль угловой двери",
            "L = G + 23; qty = 2 × q",
            "Дверь у угла",
            "ТЗ\nПРЕДВАРИТЕЛЬНО",
        ],
    ]


def hardware_rows() -> list[list[str]]:
    return [
        ["RBP0004 · скользящее покрытие", "ceil(W × 2 / 1000) × q м.п.", "Все прямые секции", "1", "excel\nПОДТВЕРЖДЕНО"],
        ["RU004 · щеточный уплотнитель", "ceil(W × 2 / 1000) × q м.п.", "Все прямые секции", "1", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0009 · болт и гайка", "(Pфиз + 1 при Pфиз > 4) × q", "По панели + общий сбор > 4; перекрывает Excel", "1", "tz\nПОДТВЕРЖДЕНО"],
        ["RBA0035 · упрочнитель углов", "2 × A × q", "Сверху и снизу каждого угла", "1", "tz\nПОДТВЕРЖДЕНО"],
        ["RBA0045 · соединитель прямой", "2 × Jпрям × q", "Стыки направляющей > 6000 мм", "1", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0036 · угол 90°", "J90 × q", "Угол ровно 90°", "1", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0006 · шарнир", "D × 3 × q", "Крайние двери", "1", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0055 · заглушка направляющей", "2 × q", "Нижняя направляющая; в Excel «будут или нет»", "1", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0040 · h-уплотнитель 10 мм", "M × q", "Обычные подвижные панели", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0041 · уплотнитель с фетром", "(2 + Jугл 90…135°) × q", "Края + угловые стыки; замена 7×15 открыта", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0001 · верхний поворотный", "M × q", "Обычные подвижные панели", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0002 · нижний поворотный", "M × q", "Обычные подвижные панели", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0003 · верхняя направляющая", "M × q", "Обычные подвижные панели", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0004 · нижняя направляющая", "M × q", "Обычные подвижные панели", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0005 · ось двери", "D × 2 × q", "Крайние двери", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0010 · поворотный механизм", "2 × ceil(M / 2) × q", "По два на пару подвижных панелей", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0050 · выход направляющей", "D × q", "Крайние двери", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0211 · доводчик левый", "(Dлев внутрь + Dправ наружу) × q", "По стороне и визуальному открыванию; старый", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0212 · доводчик правый", "(Dправ внутрь + Dлев наружу) × q", "По стороне и визуальному открыванию; старый", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0219 · верхняя защёлка", "Dбез замка × q", "Двери со стеклянной ручкой; старый", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0021 · нижняя защёлка", "Dбез замка × q", "Двери со стеклянной ручкой", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0013 · шуруп 3,5×9,5", "(Pмех×2 + D×3 + Jпрям×16 + A×16) × q", "Механизмы, двери и соединения", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0013? · шуруп 3×16", "(D×2 + Dдоп×2) × q", "Повторный артикул в Excel неизвестен", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBA0026 · стеклянная ручка", "Dбез замка × q", "Двери с ручкой", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0052 · замок с ручкой", "(Dзамок + Dдоп) × q", "Крайние и дополнительные двери", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0008 · фиксатор панелей", "D × q", "Крайние двери", "2", "excel\nПОДТВЕРЖДЕНО"],
        ["RBA0014 · крышка доводчика", "Dдоп × q", "Дополнительная двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0001 · упор 90° левый", "Dдоп по стороне/открыванию × q", "Условие стороны неоднозначно", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0002 · упор 90° правый", "Dдоп по стороне/открыванию × q", "Условие стороны неоднозначно", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM00021 · клепка-заглушка", "RBM0001 + RBM0002", "Без повторного умножения на q", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0003 · шпонка", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0004 · верхний шарнир", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0005 · нижний шарнир", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0006 · верхний поворотный", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0007 · нижний поворотный", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0008 · нижняя направляющая", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0009 · выход направляющей", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
        ["RBM0011 · контроль-деталь", "Dдоп × q", "Доп. двигающаяся дверь", "2", "excel\nПРЕДВАРИТЕЛЬНО"],
    ]


def opening_rows() -> list[list[str]]:
    return [
        ["1", "Угловые конструкции", "Подтвердить распределение ширин, допустимые углы, профиль двери +23 и состав стыка.", "Блокирует документы"],
        ["2", "Дополнительная дверь", "Заполнить пустой блок Excel: геометрия, ограничения положения, доводчик и упоры.", "Блокирует документы"],
        ["3", "Дополнительная глухая панель", "Подтвердить, входит ли она в P и как уменьшает доступный пролёт.", "Блокирует документы"],
        ["4", "Высотные вычеты", "Подтвердить H − 135 и Hстекла + 33 для всех вариантов компенсатора.", "Влияет на стекло"],
        ["5", "Отверстия и дренаж", "Уточнить базу размеров D13/D6, крайние остатки и правила лево/право/центр.", "Не реализовано"],
        ["6", "Артикулы и единицы", "RBA0055, повтор RBA0013, старые RBA0211/12/19, RBM00021; шт. или комплекты.", "Влияет на заказ"],
        ["7", "Фетровый уплотнитель", "Подтвердить замену RBA0041 на фетр 7×15 и формулу углов 90…135°.", "Влияет на заказ"],
        ["8", "Производственное округление", "Согласовать округление 0,1 мм до целого отдельно для стекла и профилей.", "Следующий пакет"],
    ]


def control_example_rows() -> list[list[str]]:
    rows = []
    width = 3000.0
    for panels in range(2, 7):
        glass = round((width - 11.5 - 11.5 - 3 * (panels - 1)) / panels, 1)
        rows.append(
            [
                str(panels),
                "3000,0",
                f"{glass:.1f}".replace(".", ","),
                f"{glass + 3:.1f}".replace(".", ","),
                "tz · ПОДТВЕРЖДЕНО",
            ]
        )
    return rows


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_opening(doc)

    add_heading(doc, "1. Источники и правило приоритета", 1)
    add_paragraph(
        doc,
        "Реестр фиксирует не только значение, но и область применения, источник "
        "и статус каждой формулы. Статус конфигурации и статус отдельных "
        "размеров хранятся раздельно.",
    )
    add_paragraph(doc, "Источник: комплект файлов для этапа 1.", style="Table Citation")
    add_table(
        doc,
        ["Приоритет", "Источник", "Что берём", "Условие применения"],
        source_rows(),
        [900, 2500, 3400, 2560],
        font_size=8.7,
        center_cols={0},
    )
    add_callout(
        doc,
        "Зафиксированный конфликт ширины",
        "Excel содержит ROUNDDOWN((W − 20 − 3P − 0,5(P−1)) / P; 0). "
        "Для прямой секции эта формула не используется, потому что ТЗ даёт "
        "собственную подтверждённую формулу.",
        LIGHT_RED,
    )

    add_heading(doc, "2. Переменные расчетного контура", 1)
    add_table(
        doc,
        ["Переменная", "Смысл", "Ед.", "Ограничение / примечание"],
        variable_rows(),
        [1200, 3500, 900, 3760],
        mono_cols={0},
        center_cols={0, 2},
    )

    add_heading(doc, "3. Геометрия физических панелей", 1)
    add_formula(doc, "G = (W − 11,5 − 11,5 − 3 × (P − 1)) / P")
    add_formula(doc, "Gp = G + 3")
    add_paragraph(
        doc,
        "Результат строится вокруг физических панелей. Для каждой панели "
        "сохраняются номер, положение, роль, направление движения, дверь и "
        "открывание, стекло, профиль, размеры, комплектующие и источник размера.",
    )
    add_table(
        doc,
        ["Расчёт", "Формула", "Область применения", "Источник", "Статус"],
        geometry_rows(),
        [1600, 2200, 2500, 1260, 1800],
        status_col=4,
        mono_cols={1},
        font_size=8.1,
    )
    add_callout(
        doc,
        "Правило ошибки",
        "Нулевой или отрицательный размер не передаётся в результат. API "
        "возвращает понятную ошибку с указанием параметров, которые следует проверить.",
        LIGHT_GREEN,
    )

    add_heading(doc, "4. Двери, движение и компенсаторы", 1)
    add_table(
        doc,
        ["Параметр", "Допустимые значения", "Привязка", "Статус"],
        [
            ["Крайние двери", "слева / справа / с двух сторон / без дверей", "Первая и/или последняя складная физическая панель", "ПОДТВЕРЖДЕНО"],
            ["Фурнитура двери", "стеклянная ручка / замок с ручкой", "Отдельно для левой и правой двери", "ПОДТВЕРЖДЕНО"],
            ["Открывание", "изнутри внутрь / изнутри наружу / снаружи наружу / снаружи внутрь", "Отдельно для каждой двери", "ПОДТВЕРЖДЕНО"],
            ["Движение", "влево / вправо / неподвижна", "Физическая панель; при двух сборах — по числу панелей слева", "ПОДТВЕРЖДЕНО"],
            ["Компенсатор", "нижний / верхний / оба", "RBP003 и схема", "ПОДТВЕРЖДЕНО"],
            ["До препятствия", "≥ 0 мм", "Размер вида сверху", "ПОДТВЕРЖДЕНО"],
            ["Высота ручки", "0…H мм", "Физическая дверь", "ПОДТВЕРЖДЕНО"],
        ],
        [1700, 2860, 3200, 1600],
        status_col=3,
        font_size=8.2,
    )

    doc.add_page_break()
    add_heading(doc, "5. Профили", 1)
    add_table(
        doc,
        ["Артикул", "Позиция", "Формула", "Область", "Источник / статус"],
        profile_rows(),
        [1100, 1900, 2350, 2350, 1660],
        status_col=4,
        mono_cols={0, 2},
        font_size=8.2,
    )

    add_heading(doc, "6. Фурнитура и этапы отгрузки", 1)
    add_paragraph(
        doc,
        "Формулы перенесены из текстового описания листа «фурнитура TODO с "
        "формулами». Этап 1 — производство; этап 2 — комплектование/отгрузка. "
        "Нулевые позиции остаются в расчётном реестре с included=false.",
    )
    add_table(
        doc,
        ["Артикул / позиция", "Формула количества", "Условие / область", "Этап", "Источник / статус"],
        hardware_rows(),
        [2050, 2350, 2700, 600, 1660],
        status_col=4,
        mono_cols={1},
        center_cols={3},
        font_size=7.7,
    )

    add_heading(doc, "7. Отверстия и дренаж — данные ТЗ, пока без выпуска", 1)
    add_table(
        doc,
        ["Элемент", "Размеры из ТЗ", "Текущее применение", "Статус"],
        [
            ["Отверстия D13", "93; 185; край стекла −38,5; остаток −28,5", "Хранятся в реестре вопросов, в детали не выпускаются", "ПРЕДВАРИТЕЛЬНО"],
            ["Отверстия D6", "124; 246; от D13 −43; остаток −40", "Хранятся в реестре вопросов, в детали не выпускаются", "ПРЕДВАРИТЕЛЬНО"],
            ["Дренаж", "Под D6 слева/справа/по центру; для D13 упомянуто 277", "Нет полной базы размеров и условий", "ПРЕДВАРИТЕЛЬНО"],
            ["Вырез доводчика", "34; первый доводчик на 25 от стекольного профиля", "Не формирует производственный документ этапа 1", "ПРЕДВАРИТЕЛЬНО"],
            ["Выход направляющей", "11 от дверного профиля", "Учтён как позиция RBA0050; отверстие не выпускается", "ПРЕДВАРИТЕЛЬНО"],
        ],
        [1700, 2600, 3300, 1760],
        status_col=3,
        font_size=8.2,
    )

    add_heading(doc, "8. Статусы и блокировка документов", 1)
    add_table(
        doc,
        ["Сценарий", "Статус конфигурации", "Расчёт", "Производственные документы"],
        [
            ["Прямая секция, 2…6 панелей", "confirmed", "Ширина — confirmed; высота — preliminary", "Следующий пакет"],
            ["Угол слева или справа", "preliminary", "Доступен с предупреждением", "ЗАБЛОКИРОВАНО"],
            ["Доп. глухая панель", "preliminary", "Доступен с предупреждением", "ЗАБЛОКИРОВАНО"],
            ["Доп. двигающаяся дверь", "preliminary", "Доступен с предупреждением", "ЗАБЛОКИРОВАНО"],
        ],
        [2800, 1900, 2600, 2060],
        status_col=3,
        font_size=8.3,
    )

    add_heading(doc, "9. Контрольные примеры", 1)
    add_paragraph(
        doc,
        "Исходные параметры для примеров: W = 3000,0 мм; H = 2500,0 мм; "
        "прямая секция; без дополнительных элементов. Ожидаемое значение "
        "фиксируется с точностью 0,1 мм.",
    )
    add_table(
        doc,
        ["P", "W, мм", "G, мм", "Gp, мм", "Источник / статус"],
        control_example_rows(),
        [900, 1700, 1800, 1800, 3160],
        status_col=4,
        mono_cols={0, 1, 2, 3},
        center_cols={0, 1, 2, 3},
        font_size=8.8,
    )
    add_paragraph(
        doc,
        "Контроль 4 панелей: (3000 − 11,5 − 11,5 − 3×3) / 4 = 742,0 мм; "
        "профиль = 745,0 мм. Excel-конфликт 741 мм намеренно не применяется.",
        bold_prefix="Контроль 4 панелей:",
    )

    add_heading(doc, "10. Открытые вопросы для согласования", 1)
    add_table(
        doc,
        ["№", "Тема", "Что требуется подтвердить", "Влияние"],
        opening_rows(),
        [600, 2000, 4700, 2060],
        center_cols={0},
        font_size=8.4,
    )

    doc.add_page_break()
    add_heading(doc, "11. Граница первого этапа", 1)
    add_table(
        doc,
        ["Входит сейчас", "Следующий пакет"],
        [
            [
                "Форма КНИЖКИ; отдельный backend-модуль; физические панели; "
                "стекло, профили и фурнитура; источники/статусы; две схемы; "
                "гостевой и авторизованный API; сохранение, копирование, шаблоны.",
                "Производственный лист, заказ стекла, заявка на покраску, "
                "накладная и производственное округление после согласования.",
            ]
        ],
        [4680, 4680],
        font_size=9,
    )

    add_heading(doc, "12. Решение, необходимое для продолжения", 1)
    add_callout(
        doc,
        "Согласовать калькулятор",
        "Подтвердить формулу прямой ширины, высотные вычеты и спорные позиции "
        "фурнитуры; отдельно дать исходные правила для углов, дополнительной "
        "двери, глухой панели, отверстий и дренажа. После этого можно фиксировать "
        "производственное округление и выпускать пакет документов.",
        LIGHT_YELLOW,
        trailing_space=False,
    )
    return doc


def audit_docx(path: Path) -> None:
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        header_xml = archive.read("word/header1.xml").decode("utf-8")
        footer_xml = archive.read("word/footer1.xml").decode("utf-8")

    assert 'w:pgSz w:w="12240" w:h="15840"' in document_xml
    assert 'w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"' in document_xml
    assert 'w:header="708"' in document_xml and 'w:footer="708"' in document_xml
    assert f'w:tblW w:type="dxa" w:w="{TABLE_WIDTH_DXA}"' in document_xml
    assert f'w:tblInd w:w="{TABLE_INDENT_DXA}" w:type="dxa"' in document_xml
    assert 'w:after="120"' in styles_xml and 'w:line="300"' in styles_xml
    assert 'w:before="360"' in styles_xml and 'w:after="200"' in styles_xml
    assert "КНИЖКА" in header_xml
    assert "PAGE" in footer_xml
    assert "<w:updateFields" in settings_xml


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    document.core_properties.title = "КНИЖКА — единый реестр формул"
    document.core_properties.subject = "Первый этап калькулятора КНИЖКА"
    document.core_properties.author = "Raluma"
    document.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
