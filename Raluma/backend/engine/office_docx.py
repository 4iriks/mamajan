"""Editable Word exports for production sheets and project documents."""

from __future__ import annotations

import io
from decimal import Decimal, ROUND_HALF_UP
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image

from engine.office_common import (
    BLACK,
    BRAND_DARK,
    HEADER_GRAY,
    RED,
    WHITE,
    drawing_image_streams_for_sections,
    format_dimension,
    format_number,
    image_stream,
    load_overrides,
    override_value,
)
from engine.office_diagrams import section_diagrams
from engine.office_section_data import (
    hardware_rows,
    profile_rows,
    section_summary_rows,
)
from engine.pdf import section_extra_components
from engine.project_documents import build_project_document_context
from engine.quote_diagrams import render_quote_room_png, render_quote_top_png


CHECKLIST_ROWS = [
    ("1", "Нарезка профиля по ТЗ", ""),
    ("2", "Фрезеровка профиля-замка под защелку (при наличии)", "Проверить сторону"),
    ("2", "Фрезеровка пазов в П-профиле и профиле-замке (при наличии)", ""),
    (
        "3",
        "Рассверловка боковых пристеночных профилей",
        "Под низкий порог снизу не сверлится",
    ),
    ("3", "Рассверловка порога", ""),
    ("4", "Установка роликов и заглушек на стекольный профиль", ""),
    ("4", "Установка фетрового уплотнения в верхний направляющий профиль", "7×6"),
    (
        "4",
        "Установка фетрового уплотнения в профиль-ручку",
        "7×6 + приклеить фетр к профилю",
    ),
    (
        "4",
        "Установка фетрового уплотнения в межстворочный профиль",
        "7×12 + приклеить фетр к профилю",
    ),
    ("5", "Склеить панели по ТЗ", ""),
    ("6", "Сборка фурнитуры по ТЗ", ""),
    ("6", "Сборка дополнительной фурнитуры по накладной", ""),
    ("7", "Оклейка всех профилей пленкой «Ралюма»", ""),
    ("7", "Упаковка панелей в пузырьковую пленку и картон", ""),
]


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(2)
    section.bottom_margin = Mm(2)
    section.left_margin = Mm(4)
    section.right_margin = Mm(4)


def _set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(8)
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_mm: float) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(int(width_mm * 56.6929)))
    width.set(qn("w:type"), "dxa")


def _hide_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def _set_cell_text(
    cell,
    value: Any,
    *,
    bold: bool = False,
    size: float = 8,
    color: str = BLACK,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(value if value is not None else ""))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _style_table(table, header_rows: int = 1) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        _prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if row_index < header_rows:
            _set_repeat_header(row)
            for cell in row.cells:
                _set_cell_shading(cell, HEADER_GRAY)
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _set_table_geometry_mm(table, widths_mm: tuple[float, ...]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(int(sum(widths_mm) * 56.6929)))
    width.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, cell_width in zip(row.cells, widths_mm, strict=False):
            _set_cell_width(cell, cell_width)


def _configure_document(document: Document, *, landscape: bool) -> None:
    if landscape:
        _set_landscape(document.sections[0])
    else:
        _set_portrait(document.sections[0])
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)


def _add_bar(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_text(cell, text.upper(), bold=True, size=7.5, color=WHITE)
    _set_cell_shading(cell, BRAND_DARK)


def _add_header(
    document: Document, project: object, section: object, label: str
) -> None:
    table = document.add_table(rows=1, cols=3)
    values = (
        f"ПРОЕКТ № {getattr(project, 'number', '')}",
        str(getattr(section, "name", "") or "Секция"),
        label,
    )
    for index, value in enumerate(values):
        _set_cell_text(
            table.cell(0, index),
            value,
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    table.style = "Table Grid"


def _add_picture(cell, data: bytes | io.BytesIO | None, width_mm: float) -> None:
    if not data:
        return
    stream = data if isinstance(data, io.BytesIO) else io.BytesIO(data)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    try:
        paragraph.add_run().add_picture(stream, width=Mm(width_mm))
    except Exception:
        return


def _add_picture_fitted(
    cell,
    data: bytes | io.BytesIO | None,
    *,
    max_width_mm: float,
    max_height_mm: float,
) -> None:
    if not data:
        return
    payload = data.getvalue() if isinstance(data, io.BytesIO) else data
    stream = io.BytesIO(payload)
    try:
        with Image.open(stream) as source:
            width_px, height_px = source.size
        stream.seek(0)
        paragraph = (
            cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        if width_px / max(height_px, 1) >= max_width_mm / max_height_mm:
            paragraph.add_run().add_picture(stream, width=Mm(max_width_mm))
        else:
            paragraph.add_run().add_picture(stream, height=Mm(max_height_mm))
    except Exception:
        return


def _add_summary(document: Document, section: object, calc: object) -> None:
    rows = section_summary_rows(section, calc)
    columns = 2
    table = document.add_table(rows=ceil_div(len(rows), columns), cols=4)
    for index, (label, value) in enumerate(rows):
        row = index // columns
        pair = index % columns
        _set_cell_text(table.cell(row, pair * 2), label, bold=True, size=6.5)
        _set_cell_text(table.cell(row, pair * 2 + 1), value, size=7)
        _set_cell_shading(table.cell(row, pair * 2), HEADER_GRAY)
    table.style = "Table Grid"


def _add_slide_parameters(
    document: Document,
    section: object,
    calc: object,
    overrides: dict[str, Any],
) -> None:
    """Mirror the compact parameters/comments block from the PDF sheet."""
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    left, right = table.rows[0].cells
    _set_cell_margins(left, top=35, start=70, bottom=35, end=70)
    _set_cell_margins(right, top=35, start=70, bottom=35, end=70)

    left.text = ""
    labels = (
        ("ЦВЕТ", getattr(calc, "color_text", "") or ""),
        ("СТЕКЛО", getattr(calc, "glass_type", "") or ""),
        ("ПОРОГ", getattr(calc, "threshold_text", "") or ""),
        ("СИСТЕМА", getattr(calc, "system_text", "") or ""),
    )
    for label, value in labels:
        paragraph = (
            left.add_paragraph() if left.paragraphs[0].text else left.paragraphs[0]
        )
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(f"{label:<10} ")
        label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(6.7)
        value_run = paragraph.add_run(str(value))
        value_run.font.name = "Arial"
        value_run.font.size = Pt(6.7)

    glass_table = left.add_table(rows=1, cols=4)
    for index, header in enumerate(("СТЕКЛА", "ширина", "высота", "кол-во")):
        _set_cell_text(
            glass_table.cell(0, index),
            header,
            bold=True,
            size=5.8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for index, glass in enumerate(getattr(calc, "glass", None) or []):
        if int(getattr(glass, "qty", 0) or 0) <= 0:
            continue
        row = glass_table.add_row()
        values = (
            glass.position,
            override_value(
                overrides,
                f"glass_{index}_w",
                format_dimension(glass.width_mm),
            ),
            override_value(
                overrides,
                f"glass_{index}_h",
                format_dimension(glass.height_mm),
            ),
            override_value(overrides, f"glass_{index}_q", glass.qty),
        )
        for column, value in enumerate(values):
            _set_cell_text(
                row.cells[column],
                value,
                size=5.8,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
    _style_table(glass_table)

    right.text = ""
    inner = right.add_table(rows=1, cols=2)
    metrics_cell, comments_cell = inner.rows[0].cells
    metrics = (
        ("ШИРИНА СЕКЦИИ, мм", getattr(section, "width", 0)),
        ("ВЫСОТА СЕКЦИИ, мм", getattr(section, "height", 0)),
        ("КОЛИЧЕСТВО ПАНЕЛЕЙ, шт", getattr(section, "panels", 0)),
        ("КОЛИЧЕСТВО СЕКЦИЙ, шт", getattr(section, "quantity", 0)),
    )
    metrics_cell.text = ""
    for label, value in metrics:
        paragraph = (
            metrics_cell.add_paragraph()
            if metrics_cell.paragraphs[0].text
            else metrics_cell.paragraphs[0]
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(f"{label}  ")
        label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(6.1)
        value_run = paragraph.add_run(
            format_dimension(value) if "СЕКЦИИ, мм" in label else str(value)
        )
        value_run.bold = True
        value_run.font.name = "Arial"
        value_run.font.size = Pt(6.3)

    comments_cell.text = ""
    title = comments_cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ПРИМЕЧАНИЕ")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(6.5)
    comments = override_value(
        overrides,
        "section_comments",
        getattr(section, "comments", "") or "",
    )
    content = comments_cell.add_paragraph(str(comments))
    content.paragraph_format.space_before = Pt(1)
    content.paragraph_format.space_after = Pt(0)
    for run in content.runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(7)
    inner.style = "Table Grid"


def ceil_div(value: int, divisor: int) -> int:
    return (value + divisor - 1) // divisor


def _add_diagrams(document: Document, section: object, calc: object) -> None:
    diagrams = section_diagrams(section, calc)
    first = diagrams[:2]
    is_lift = str(getattr(section, "system", "") or "").strip().upper() == "ЛИФТ"
    table = document.add_table(rows=1, cols=len(first))
    table.autofit = False
    picture_width = 78 if is_lift else 90
    cell_width = 96 if not is_lift else 88
    placements = []
    for index in range(len(first)):
        cell = table.cell(0, index)
        _set_cell_width(cell, cell_width)
        _set_cell_margins(cell, top=20, start=20, bottom=20, end=20)
        placements.append((cell, picture_width))

    for (title, data), (cell, width_mm) in zip(first, placements, strict=True):
        _set_cell_text(
            cell,
            title,
            bold=True,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run().add_picture(
            io.BytesIO(data),
            width=Mm(width_mm),
        )
    table.style = "Table Grid"

    for title, data in diagrams[2:]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(
            io.BytesIO(data),
            width=Mm(174),
        )


def _add_calc_warnings(document: Document, calc: object) -> None:
    for warning_text in getattr(calc, "warnings", []) or []:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(str(warning_text))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string(RED)


def _add_slide_glass(
    document: Document, calc: object, overrides: dict[str, Any]
) -> None:
    _add_bar(document, "Стекла")
    table = document.add_table(rows=1, cols=5)
    headers = ("Позиция", "Ширина, мм", "Высота, мм", "Кол-во, шт", "RS2021, мм")
    for index, header in enumerate(headers):
        _set_cell_text(
            table.cell(0, index),
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for index, glass in enumerate(getattr(calc, "glass", None) or []):
        if int(getattr(glass, "qty", 0) or 0) <= 0:
            continue
        row = table.add_row()
        values = (
            glass.position,
            override_value(
                overrides,
                f"glass_{index}_w",
                format_dimension(glass.width_mm),
            ),
            override_value(
                overrides,
                f"glass_{index}_h",
                format_dimension(glass.height_mm),
            ),
            override_value(overrides, f"glass_{index}_q", glass.qty),
            format_dimension(glass.glass_profile_length),
        )
        for column, value in enumerate(values):
            _set_cell_text(
                row.cells[column],
                value,
                align=WD_ALIGN_PARAGRAPH.CENTER if column else None,
            )
    _style_table(table)


def _add_lift_panels(
    document: Document, calc: object, overrides: dict[str, Any]
) -> None:
    _add_bar(document, "Заполнение и панели при склейке")
    table = document.add_table(rows=1, cols=7)
    headers = (
        "№",
        "Панель",
        "Заполнение",
        "Ширина, мм",
        "Высота, мм",
        "Склейка, мм",
        "Кол-во",
    )
    for index, header in enumerate(headers):
        _set_cell_text(
            table.cell(0, index),
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for panel in getattr(calc, "panels", None) or []:
        row = table.add_row()
        values = (
            panel.panel,
            panel.role,
            panel.filling,
            override_value(
                overrides,
                f"lift_panel_{panel.panel}_width",
                format_dimension(panel.width_mm),
            ),
            override_value(
                overrides,
                f"lift_panel_{panel.panel}_height",
                format_dimension(panel.height_mm),
            ),
            override_value(
                overrides,
                f"lift_panel_{panel.panel}_glued_table",
                f"{format_dimension(panel.glued_width_mm)} × "
                f"{format_dimension(panel.glued_height_mm)}",
            ),
            override_value(
                overrides,
                f"lift_panel_{panel.panel}_qty",
                panel.qty,
            ),
        )
        for column, value in enumerate(values):
            _set_cell_text(
                row.cells[column],
                value,
                align=WD_ALIGN_PARAGRAPH.CENTER if column != 2 else None,
            )
    _style_table(table)


def _add_compact_cards(
    document: Document,
    title: str,
    cards: list[tuple[io.BytesIO | None, str, str, str]],
) -> None:
    if not cards:
        return
    _add_bar(document, title)
    columns = 3
    table = document.add_table(rows=ceil_div(len(cards), columns), cols=columns)
    table.style = "Table Grid"
    table.autofit = False
    for index, (image, heading, details, note) in enumerate(cards):
        cell = table.cell(index // columns, index % columns)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell, top=12, start=20, bottom=12, end=20)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.add_run().font.size = Pt(1)

        nested = cell.add_table(rows=2 if note else 1, cols=3)
        nested.autofit = False
        _hide_table_borders(nested)
        for nested_cell, width in zip(
            nested.rows[0].cells,
            (15, 34, 14),
            strict=True,
        ):
            _set_cell_width(nested_cell, width)
            _set_cell_margins(nested_cell, top=0, start=25, bottom=0, end=25)
            nested_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        if image:
            try:
                image.seek(0)
                with Image.open(image) as source:
                    scale = min(13.5 / source.width, 8.5 / source.height)
                    width_mm = max(1, source.width * scale)
                    height_mm = max(1, source.height * scale)
                image.seek(0)
                image_paragraph = nested.cell(0, 0).paragraphs[0]
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.paragraph_format.space_before = Pt(0)
                image_paragraph.paragraph_format.space_after = Pt(0)
                image_paragraph.add_run().add_picture(
                    image,
                    width=Mm(width_mm),
                    height=Mm(height_mm),
                )
            except Exception:
                pass

        heading_paragraph = nested.cell(0, 1).paragraphs[0]
        heading_paragraph.paragraph_format.space_before = Pt(0)
        heading_paragraph.paragraph_format.space_after = Pt(0)
        heading_run = heading_paragraph.add_run(heading)
        heading_run.bold = True
        heading_run.font.name = "Arial"
        heading_run.font.size = Pt(7.2)

        if details:
            details_paragraph = nested.cell(0, 2).paragraphs[0]
            details_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            details_paragraph.paragraph_format.space_before = Pt(0)
            details_paragraph.paragraph_format.space_after = Pt(0)
            details_run = details_paragraph.add_run(details)
            details_run.bold = True
            details_run.font.name = "Arial"
            details_run.font.size = Pt(8.2)

        if note:
            note_cell = nested.cell(1, 0).merge(nested.cell(1, 2))
            _set_cell_margins(note_cell, top=0, start=25, bottom=0, end=25)
            note_paragraph = note_cell.paragraphs[0]
            note_paragraph.paragraph_format.space_before = Pt(0)
            note_paragraph.paragraph_format.space_after = Pt(0)
            note_run = note_paragraph.add_run(note)
            note_run.italic = True
            note_run.font.name = "Arial"
            note_run.font.size = Pt(7)
            note_run.font.color.rgb = RGBColor.from_string("555555")
    for row_index, row in enumerate(table.rows):
        _prevent_row_split(row)
        start = row_index * columns
        has_note = any(card[3] for card in cards[start : start + columns])
        row.height = Mm(16 if has_note else 13)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _add_profiles(document: Document, calc: object, overrides: dict[str, Any]) -> None:
    cards: list[tuple[io.BytesIO | None, str, str, str]] = []
    for index, profile in enumerate(profile_rows(calc)):
        cuts = list(getattr(profile, "display_cuts", None) or [])
        if not cuts:
            cuts = [
                {
                    "length": getattr(profile, "length_mm", 0),
                    "qty": getattr(profile, "qty", 0),
                    "length_field": getattr(profile, "field_key", "")
                    or f"profile_{index}_length",
                    "qty_field": (
                        getattr(profile, "field_key", "") or f"profile_{index}"
                    )
                    + "_qty",
                }
            ]
        cut_labels: list[str] = []
        for cut in cuts:
            length = ""
            if str(getattr(profile, "article", "")).upper() not in {
                "RS1005",
                "RS3110",
            }:
                length = override_value(
                    overrides,
                    str(cut.get("length_field") or ""),
                    format_dimension(cut.get("length")),
                )
            qty = override_value(
                overrides,
                str(cut.get("qty_field") or ""),
                cut.get("qty", 0),
            )
            if length == "":
                cut_labels.append(f"{qty} шт")
            else:
                cut_labels.append(f"{length} мм {qty} шт")
        article = str(getattr(profile, "article", "") or "")
        name = str(getattr(profile, "name", "") or "")
        cards.append(
            (
                image_stream(getattr(profile, "image", None), max_size=(600, 260)),
                f"{article} · {name}",
                "; ".join(cut_labels),
                str(getattr(profile, "note", "") or ""),
            )
        )
    _add_compact_cards(document, "Нарезка профилей", cards)


def _add_hardware(document: Document, calc: object, overrides: dict[str, Any]) -> None:
    cards: list[tuple[io.BytesIO | None, str, str, str]] = []
    for article, name, value, unit, image, field_key, note in hardware_rows(calc):
        quantity = override_value(overrides, field_key, format_number(value))
        cards.append(
            (
                image_stream(image, max_size=(500, 260)),
                f"{article} · {name}",
                f"{quantity} {unit}".strip(),
                str(note or ""),
            )
        )
    _add_compact_cards(document, "Фурнитура и крепеж", cards)


def _add_extra_components(
    document: Document, section: object, overrides: dict[str, Any]
) -> None:
    rows = section_extra_components(section, overrides)
    if not rows:
        return
    _add_bar(document, "Дополнительные комплектующие")
    table = document.add_table(rows=1, cols=5)
    headers = ("Артикул", "Название", "Размер", "Кол-во", "Цвет")
    for index, header in enumerate(headers):
        _set_cell_text(table.cell(0, index), header, bold=True, size=6.5)
    for item in rows:
        row = table.add_row()
        for index, key in enumerate(("art", "name", "size", "qty", "color")):
            _set_cell_text(row.cells[index], item.get(key, ""), size=6)
    _style_table(table)


def _add_slide_checklist(
    document: Document,
    project: object,
    section: object,
    overrides: dict[str, Any],
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        f"ПРОЕКТ № {getattr(project, 'number', '')} — {getattr(section, 'name', '')}"
    )
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    table = document.add_table(rows=1, cols=5)
    table.autofit = False
    column_widths = (13, 19, 78, 64, 19)
    headers = ("№ п/п", "Отм. пр-ва", "Действие", "Примечание", "Отм. ОТК")
    for index, header in enumerate(headers):
        _set_cell_width(table.cell(0, index), column_widths[index])
        _set_cell_text(
            table.cell(0, index),
            header,
            bold=True,
            size=7.4,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    table.rows[0].height = Mm(8)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for number, action, note in CHECKLIST_ROWS:
        row = table.add_row()
        values = (number, "☐", action, note, "☐")
        for index, value in enumerate(values):
            _set_cell_width(row.cells[index], column_widths[index])
            _set_cell_text(
                row.cells[index],
                value,
                size=7.2,
                align=WD_ALIGN_PARAGRAPH.CENTER if index in {0, 1, 4} else None,
            )
        row.height = Mm(8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _style_table(table)
    _add_bar(document, "Примечания и особые отметки при производстве или проверке ОТК")
    comments_table = document.add_table(rows=1, cols=1)
    comments = override_value(
        overrides, "section_comments", getattr(section, "comments", "") or ""
    )
    _set_cell_text(comments_table.cell(0, 0), comments, bold=True, size=9)
    comments_table.rows[0].height = Mm(29)
    comments_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    comments_table.style = "Table Grid"

    _add_bar(document, "Ответственные за заказ на производстве")
    people = document.add_table(rows=4, cols=6)
    labels = (
        "Нарезка",
        "Поклейка",
        "Упаковка",
        "Сборка",
        "Поклейка",
        "Комплектация",
        "Упаковка",
        "",
    )
    for index, label in enumerate(labels):
        row, pair = divmod(index, 2)
        start = pair * 3
        _set_cell_text(
            people.cell(row, start),
            index + 1,
            size=7.3,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text(people.cell(row, start + 1), label, size=7.3)
        _set_cell_text(people.cell(row, start + 2), " - ____________________", size=7.3)
        people.rows[row].height = Mm(6.2)
        people.rows[row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    people.style = "Table Grid"

    footer = document.add_table(rows=5, cols=2)
    footer.style = "Table Grid"
    footer.cell(0, 0).merge(footer.cell(1, 0))
    footer.cell(2, 0).merge(footer.cell(4, 0))
    footer.cell(0, 1).merge(footer.cell(0, 1))
    _set_cell_text(
        footer.cell(0, 0),
        'Дата фото профиля и панелей   "____"  ______________  202___ г.',
        size=8,
    )
    _set_cell_text(
        footer.cell(2, 0),
        'Дата фото фурнитуры           "____"  ______________  202___ г.',
        size=8,
    )
    _set_cell_text(
        footer.cell(0, 1),
        "Ответственный",
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        footer.cell(1, 1),
        "________________________________",
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        footer.cell(2, 1),
        "подпись",
        size=7,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        footer.cell(3, 1),
        "________________________________",
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        footer.cell(4, 1),
        "расшифровка",
        size=7,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for footer_row in footer.rows:
        footer_row.height = Mm(5.2)
        footer_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _add_slide_drawings(document: Document, section: object) -> None:
    drawings = drawing_image_streams_for_sections([section])
    if not drawings:
        return
    _add_bar(document, "Чертежи ручек")
    table = document.add_table(rows=1, cols=len(drawings))
    table.autofit = False
    for cell, (name, stream) in zip(table.rows[0].cells, drawings, strict=True):
        _add_picture(cell, stream, 34)
        label = "Ручка-кноб" if name == "knob" else "Ручка-скоба 600 мм"
        paragraph = cell.add_paragraph(label)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(6.5)
    table.style = "Table Grid"


def build_section_docx(project: object, section: object, calc: object) -> bytes:
    document = Document()
    _configure_document(document, landscape=False)
    system = str(getattr(section, "system", "") or "").strip().upper()
    label = (
        "ЛИФТ · ПРОИЗВОДСТВЕННЫЙ ЛИСТ"
        if system == "ЛИФТ"
        else "СЛАЙД · ПРОИЗВОДСТВЕННЫЙ ЛИСТ"
    )
    overrides = load_overrides(section)
    _add_header(document, project, section, label)
    _add_calc_warnings(document, calc)
    if system == "ЛИФТ":
        _add_summary(document, section, calc)
        _add_diagrams(document, section, calc)
        _add_lift_panels(document, calc, overrides)
        document.add_page_break()
        _add_header(document, project, section, "ЛИФТ · НАРЕЗКА И КОМПЛЕКТАЦИЯ")
    else:
        _add_diagrams(document, section, calc)
        _add_slide_parameters(document, section, calc, overrides)
    _add_profiles(document, calc, overrides)
    _add_hardware(document, calc, overrides)
    _add_extra_components(document, section, overrides)

    if system == "ЛИФТ":
        if getattr(calc, "torque", None):
            _add_bar(document, "Расчет привода")
            torque = calc.torque
            table = document.add_table(rows=2, cols=3)
            values = (
                (
                    "Вес подвижных панелей",
                    f"{format_dimension(torque.moving_weight_kg)} кг",
                ),
                (
                    "Крутящий момент",
                    f"{format_dimension(torque.torque_nm)} Н·м",
                ),
                ("Количество приводов", f"{torque.drive_count} шт"),
            )
            for index, (label_text, value) in enumerate(values):
                _set_cell_text(
                    table.cell(0, index),
                    label_text,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                _set_cell_text(
                    table.cell(1, index),
                    value,
                    bold=True,
                    size=10,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
            table.style = "Table Grid"
        _add_bar(document, "Примечания и особые отметки")
        notes = document.add_table(rows=1, cols=1)
        _set_cell_text(
            notes.cell(0, 0),
            override_value(
                overrides, "lift_comments", getattr(section, "comments", "") or ""
            ),
            bold=True,
            size=10,
        )
        notes.rows[0].height = Mm(20)
        notes.style = "Table Grid"
    else:
        document.add_page_break()
        _add_slide_checklist(document, project, section, overrides)
        _add_slide_drawings(document, section)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _project_header(document: Document, title: str, project: object) -> None:
    table = document.add_table(rows=2, cols=3)
    title_cell = table.cell(0, 0).merge(table.cell(1, 0))
    _set_cell_text(title_cell, title.upper(), bold=True, size=16)
    _set_cell_text(table.cell(0, 1), "Заявка", bold=True, size=8)
    _set_cell_text(
        table.cell(0, 2),
        str(getattr(project, "number", "") or ""),
        size=8,
    )
    _set_cell_text(table.cell(1, 1), "Заказчик", bold=True, size=8)
    _set_cell_text(
        table.cell(1, 2),
        str(getattr(project, "customer", "") or ""),
        size=8,
    )
    table.style = "Table Grid"


def _add_project_document_warnings(document: Document, context: dict) -> None:
    for text in context.get("document_warnings", []):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(str(text))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(RED)


def _sketch_heading(document: Document, text: str, *, size: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(BRAND_DARK)


def _build_sketch_docx(context: dict) -> bytes:
    document = Document()
    _configure_document(document, landscape=False)
    section = document.sections[0]
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(11)
    section.right_margin = Mm(11)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(0)
    document.core_properties.title = context["document_title"]
    document.core_properties.subject = "Эскизный проект"
    document.core_properties.company = "Raluma"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(context["document_title"])
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor.from_string(BRAND_DARK)

    if not context["sections"]:
        paragraph = document.add_paragraph("В проекте нет секций.")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section_index, sketch_section in enumerate(context["sections"]):
        if section_index:
            document.add_page_break()
        _sketch_heading(
            document,
            f"Изделие № {sketch_section['order']} · {sketch_section['system_text']}",
            size=10.5,
        )
        for warning_text in sketch_section["warnings"]:
            warning = document.add_paragraph()
            warning.paragraph_format.space_before = Pt(1)
            warning.paragraph_format.space_after = Pt(3)
            warning.paragraph_format.keep_with_next = True
            warning_run = warning.add_run(str(warning_text))
            warning_run.bold = True
            warning_run.font.name = "Arial"
            warning_run.font.size = Pt(7.5)
            warning_run.font.color.rgb = RGBColor.from_string(RED)

        meta = document.add_table(rows=5, cols=4)
        meta_values = (
            (
                "Номер изделия",
                sketch_section["order"],
                "Система",
                sketch_section["system_text"],
            ),
            (
                "Ширина",
                f"{format_dimension(sketch_section['width_mm'])} мм",
                "Высота",
                f"{format_dimension(sketch_section['height_mm'])} мм",
            ),
            (
                "Цвет",
                sketch_section["color"],
                "Количество",
                f"{sketch_section['quantity']} шт",
            ),
            (
                "Порог",
                sketch_section["threshold"],
                "Межстекольный профиль",
                sketch_section["inter_glass_profile"],
            ),
        )
        for row_index, values in enumerate(meta_values):
            for column, value in enumerate(values):
                _set_cell_text(
                    meta.cell(row_index, column),
                    value,
                    bold=column % 2 == 0,
                    size=6.7 if column % 2 == 0 else 7.2,
                )
                if column % 2 == 0:
                    _set_cell_shading(meta.cell(row_index, column), HEADER_GRAY)
        _set_cell_text(meta.cell(4, 0), "Стекло / заполнение", bold=True, size=6.7)
        _set_cell_shading(meta.cell(4, 0), HEADER_GRAY)
        filling_cell = meta.cell(4, 1).merge(meta.cell(4, 3))
        _set_cell_text(filling_cell, sketch_section["filling"], bold=True, size=7.2)
        meta.style = "Table Grid"
        _set_table_geometry_mm(meta, (32, 59, 32, 59))
        for row in meta.rows:
            _prevent_row_split(row)

        if sketch_section.get("comments"):
            comment = document.add_paragraph()
            comment.paragraph_format.space_before = Pt(2)
            comment.paragraph_format.space_after = Pt(3)
            comment.paragraph_format.keep_with_next = True
            comment_run = comment.add_run(
                f"Примечание: {sketch_section['comments']}"
            )
            comment_run.font.name = "Arial"
            comment_run.font.size = Pt(7.2)
            comment_run.italic = True

        diagrams = sketch_section["diagrams"]
        if diagrams:
            diagram_table = document.add_table(rows=len(diagrams), cols=1)
            for index, diagram in enumerate(diagrams):
                cell = diagram_table.cell(index, 0)
                _set_cell_text(
                    cell,
                    diagram["title"].upper(),
                    bold=True,
                    size=7,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                _add_picture_fitted(
                    cell,
                    diagram["png"],
                    max_width_mm=176,
                    max_height_mm=58,
                )
            diagram_table.style = "Table Grid"
            _set_table_geometry_mm(diagram_table, (182,))
            for row in diagram_table.rows:
                _prevent_row_split(row)

        _sketch_heading(document, "РАЗМЕРЫ СТЕКОЛ", size=8)
        panels_table = document.add_table(rows=1, cols=6)
        panel_headers = (
            "№",
            "Позиция",
            "Стекло / заполнение",
            "Ширина, мм",
            "Высота, мм",
            "Кол-во",
        )
        for column, header in enumerate(panel_headers):
            _set_cell_text(
                panels_table.cell(0, column),
                header,
                bold=True,
                size=6.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        for panel in sketch_section["panels"]:
            row = panels_table.add_row()
            values = (
                panel.number,
                panel.position,
                panel.filling,
                format_dimension(panel.width_mm),
                format_dimension(panel.height_mm),
                panel.qty,
            )
            for column, value in enumerate(values):
                _set_cell_text(
                    row.cells[column],
                    value,
                    size=6.5,
                    align=(
                        WD_ALIGN_PARAGRAPH.CENTER
                        if column in {0, 3, 4, 5}
                        else WD_ALIGN_PARAGRAPH.LEFT
                    ),
                )
        _style_table(panels_table)
        _set_table_geometry_mm(panels_table, (12, 32, 70, 25, 25, 18))
        for row in panels_table.rows:
            _prevent_row_split(row)

        _sketch_heading(document, "КОМПЛЕКТАЦИЯ", size=8)
        components = sketch_section["components"]
        if components:
            components_table = document.add_table(rows=1, cols=5)
            component_headers = (
                "Артикул",
                "Наименование",
                "Размер",
                "Кол-во",
                "Примечание",
            )
            for column, header in enumerate(component_headers):
                _set_cell_text(
                    components_table.cell(0, column),
                    header,
                    bold=True,
                    size=6.5,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
            for component in components:
                row = components_table.add_row()
                component_meta = "\n".join(
                    filter(
                        None,
                        (
                            component["name"],
                            (
                                f"Цвет: {component['color']}"
                                if component.get("color")
                                else ""
                            ),
                            (
                                f"Этап: {component['stage']}"
                                if component.get("stage")
                                else ""
                            ),
                        ),
                    )
                )
                values = (
                    component["article"] or "—",
                    component_meta,
                    component["size"],
                    f"{component['qty']} {component['unit']}",
                    component["note"] or "—",
                )
                for column, value in enumerate(values):
                    _set_cell_text(
                        row.cells[column],
                        value,
                        size=6.3,
                        align=(
                            WD_ALIGN_PARAGRAPH.CENTER
                            if column in {0, 2, 3}
                            else WD_ALIGN_PARAGRAPH.LEFT
                        ),
                    )
                _add_picture_fitted(
                    row.cells[1],
                    image_stream(component.get("image"), max_size=(500, 300)),
                    max_width_mm=14,
                    max_height_mm=7,
                )
            _style_table(components_table)
            _set_table_geometry_mm(components_table, (28, 70, 28, 22, 34))
            for row in components_table.rows:
                _prevent_row_split(row)
        else:
            empty = document.add_paragraph("Комплектация не выбрана.")
            empty.paragraph_format.space_after = Pt(3)

    project_components = context.get("project_components") or []
    project_note = str(context.get("project_extra_parts_note") or "").strip()
    if project_components or project_note:
        if context["sections"]:
            document.add_page_break()
        _sketch_heading(
            document,
            "ДОПОЛНИТЕЛЬНЫЕ КОМПЛЕКТУЮЩИЕ ПРОЕКТА",
            size=10.5,
        )
        if project_note:
            note = document.add_paragraph()
            note.paragraph_format.space_after = Pt(4)
            note_run = note.add_run(f"Примечание к комплектации: {project_note}")
            note_run.font.name = "Arial"
            note_run.font.size = Pt(7.5)
            note_run.italic = True
        if project_components:
            project_table = document.add_table(rows=1, cols=5)
            headers = (
                "Артикул",
                "Наименование",
                "Размер",
                "Кол-во",
                "Примечание",
            )
            for column, header in enumerate(headers):
                _set_cell_text(
                    project_table.cell(0, column),
                    header,
                    bold=True,
                    size=6.5,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
            for component in project_components:
                row = project_table.add_row()
                component_meta = "\n".join(
                    filter(
                        None,
                        (
                            component["name"],
                            f"Цвет: {component['color']}" if component.get("color") else "",
                            f"Этап: {component['stage']}" if component.get("stage") else "",
                        ),
                    )
                )
                values = (
                    component["article"] or "—",
                    component_meta,
                    component["size"],
                    f"{component['qty']} {component['unit']}",
                    component["note"] or "—",
                )
                for column, value in enumerate(values):
                    _set_cell_text(
                        row.cells[column],
                        value,
                        size=6.3,
                        align=(
                            WD_ALIGN_PARAGRAPH.CENTER
                            if column in {0, 2, 3}
                            else WD_ALIGN_PARAGRAPH.LEFT
                        ),
                    )
                _add_picture_fitted(
                    row.cells[1],
                    image_stream(component.get("image"), max_size=(500, 300)),
                    max_width_mm=14,
                    max_height_mm=7,
                )
            _style_table(project_table)
            _set_table_geometry_mm(project_table, (28, 70, 28, 22, 34))
            for row in project_table.rows:
                _prevent_row_split(row)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _rubles(value: Any) -> str:
    try:
        rounded = int(
            Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
        )
    except Exception:
        rounded = 0
    return f"{rounded:,}".replace(",", " ")


def _build_commercial_docx(context: dict) -> bytes:
    document = Document()
    _configure_document(document, landscape=True)
    quote = context["commercial_quote"]
    _project_header(document, "Коммерческое предложение", context["project"])

    status = document.add_paragraph()
    status.paragraph_format.space_before = Pt(3)
    status.paragraph_format.space_after = Pt(3)
    status_run = status.add_run(
        f"Редакция {quote['revision']} · "
        + ("зафиксировано" if quote.get("status") == "fixed" else "черновой расчёт")
    )
    status_run.font.name = "Arial"
    status_run.font.size = Pt(8)
    status_run.font.color.rgb = RGBColor.from_string("666666")
    _add_project_document_warnings(document, context)

    construction_rows = [
        row
        for row in quote.get("lines") or []
        if isinstance(row.get("section_details"), dict)
    ]
    if construction_rows:
        technical_label = document.add_paragraph()
        technical_label.paragraph_format.space_before = Pt(3)
        technical_label.paragraph_format.space_after = Pt(2)
        technical_label_run = technical_label.add_run("ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ")
        technical_label_run.bold = True
        technical_label_run.font.name = "Arial"
        technical_label_run.font.size = Pt(9)
        technical = document.add_table(rows=1, cols=9)
        technical_headers = (
            "Изделие",
            "Габарит, мм",
            "Панели",
            "Кол-во",
            "Стекло, м²",
            "Цвет",
            "Система",
            "Тип стекла",
            "Порог",
        )
        technical_widths = (38, 25, 15, 15, 20, 29, 24, 70, 53)
        for index, header in enumerate(technical_headers):
            _set_cell_text(
                technical.cell(0, index),
                header,
                bold=True,
                size=6.4,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_width(technical.cell(0, index), technical_widths[index])
        for line in construction_rows:
            details = line["section_details"]
            row = technical.add_row()
            values = (
                line.get("name", ""),
                f"{details.get('width_mm', '')} × {details.get('height_mm', '')}",
                details.get("panels", ""),
                details.get("quantity", ""),
                details.get("glass_area_m2", ""),
                details.get("color", ""),
                f"{details.get('system', '')} · {details.get('rails', '')} р.",
                details.get("glass_type", ""),
                details.get("threshold", ""),
            )
            for index, value in enumerate(values):
                _set_cell_text(
                    row.cells[index],
                    value,
                    bold=index == 0,
                    size=6.5,
                    align=(
                        WD_ALIGN_PARAGRAPH.CENTER if index in {1, 2, 3, 4, 6} else None
                    ),
                )
                _set_cell_width(row.cells[index], technical_widths[index])
        _style_table(technical)

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    table = document.add_table(rows=1, cols=8)
    headers = (
        "№",
        "Изделие / услуга",
        "Цена до скидки, ₽",
        "Скидка, %",
        "Сумма скидки, ₽",
        "Цена, ₽",
        "Кол-во",
        "Итого, ₽",
    )
    widths = (7, 78, 34, 22, 34, 32, 30, 52)
    for index, header in enumerate(headers):
        _set_cell_text(
            table.cell(0, index),
            header,
            bold=True,
            size=6.8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_width(table.cell(0, index), widths[index])
    for row_index, row_data in enumerate(quote.get("lines") or [], start=1):
        row = table.add_row()
        values = (
            row_index,
            row_data.get("name", ""),
            _rubles(row_data.get("document_unit_price_before_discount")),
            row_data.get("discount_percent", "0"),
            _rubles(row_data.get("document_unit_discount_amount")),
            _rubles(row_data.get("document_unit_final_price")),
            f"{row_data.get('quantity', '0')} {row_data.get('unit', '')}".strip(),
            _rubles(row_data.get("document_line_total")),
        )
        for index, value in enumerate(values):
            _set_cell_text(
                row.cells[index],
                value,
                bold=index in {1, 7},
                size=7,
                align=(
                    None
                    if index == 1
                    else WD_ALIGN_PARAGRAPH.RIGHT
                    if index in {2, 4, 5, 7}
                    else WD_ALIGN_PARAGRAPH.CENTER
                ),
            )
            _set_cell_width(row.cells[index], widths[index])
    _style_table(table)

    totals = quote["totals"]
    totals_table = document.add_table(rows=0, cols=2)
    totals_rows = [
        ("До скидки", f"{_rubles(totals['document_before_discount'])} ₽"),
        ("Скидка", f"−{_rubles(totals['document_discount'])} ₽"),
    ]
    vat = quote.get("vat") or {}
    if vat.get("mode") == "included":
        totals_rows.append(
            (
                f"В том числе НДС {vat.get('rate', '0')}%",
                f"{_rubles(vat.get('document_amount', vat.get('amount', 0)))} ₽",
            )
        )
    elif vat.get("mode") == "on_top":
        totals_rows.append(
            (
                f"НДС {vat.get('rate', '0')}% сверху",
                f"{_rubles(vat.get('document_amount', vat.get('amount', 0)))} ₽",
            )
        )
    else:
        totals_rows.append(("НДС", "Без НДС"))
    totals_rows.append(("ИТОГО", f"{_rubles(totals['document_grand_total'])} ₽"))
    for row_index, (label, value) in enumerate(totals_rows):
        row = totals_table.add_row()
        _set_cell_text(
            row.cells[0],
            label,
            bold=row_index == len(totals_rows) - 1,
            size=8.5,
        )
        _set_cell_text(
            row.cells[1],
            value,
            bold=row_index == len(totals_rows) - 1,
            size=8.5,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        _set_cell_width(row.cells[0], 58)
        _set_cell_width(row.cells[1], 32)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _style_table(totals_table, header_rows=0)

    terms = document.add_paragraph()
    terms.paragraph_format.space_before = Pt(5)
    terms.paragraph_format.space_after = Pt(0)
    terms_lines = (
        f"Предложение действительно до: {quote.get('valid_until', '')}",
        f"Срок изготовления: {quote.get('manufacturing_term') or 'по согласованию'}",
        f"Условия оплаты: {quote.get('payment_terms') or 'по согласованию'}",
    )
    for index, text in enumerate(terms_lines):
        if index:
            terms.add_run("\n")
        run = terms.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(9)

    for line in construction_rows:
        details = line["section_details"]
        document.add_page_break()
        heading = document.add_paragraph()
        heading.paragraph_format.space_after = Pt(2)
        heading_run = heading.add_run(str(line.get("name") or "Секция").upper())
        heading_run.bold = True
        heading_run.font.name = "Arial"
        heading_run.font.size = Pt(14)

        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(4)
        meta_text = (
            f"{details.get('width_mm', '')} × {details.get('height_mm', '')} мм · "
            f"{details.get('panels', '')} пан. · {details.get('quantity', '')} изд. · "
            f"{details.get('system', '')}, {details.get('rails', '')} рельс. · "
            f"{details.get('glass_type', '')} · {details.get('color', '')}"
        )
        meta_run = meta.add_run(meta_text)
        meta_run.font.name = "Arial"
        meta_run.font.size = Pt(8)
        meta_run.font.color.rgb = RGBColor.from_string("555555")

        sketches = document.add_table(rows=1, cols=2)
        _hide_table_borders(sketches)
        _set_cell_width(sketches.cell(0, 0), 143)
        _set_cell_width(sketches.cell(0, 1), 143)
        _add_picture_fitted(
            sketches.cell(0, 0),
            render_quote_room_png(details),
            max_width_mm=137,
            max_height_mm=56,
        )
        _add_picture_fitted(
            sketches.cell(0, 1),
            render_quote_top_png(details),
            max_width_mm=137,
            max_height_mm=56,
        )

        label = document.add_paragraph()
        label.paragraph_format.space_before = Pt(3)
        label.paragraph_format.space_after = Pt(2)
        label_run = label.add_run("ПРОФИЛИ И ФУРНИТУРА")
        label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(9)

        breakdown = document.add_table(rows=1, cols=6)
        breakdown_headers = (
            "Артикул",
            "Название",
            "Количество",
            "Единица",
            "Продажная цена, ₽",
            "Сумма, ₽",
        )
        breakdown_widths = (28, 130, 30, 25, 36, 40)
        for index, header in enumerate(breakdown_headers):
            _set_cell_text(
                breakdown.cell(0, index),
                header,
                bold=True,
                size=6.7,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_width(breakdown.cell(0, index), breakdown_widths[index])
        for item in line.get("breakdown") or []:
            row = breakdown.add_row()
            values = (
                item.get("sku", ""),
                item.get("name", ""),
                item.get("quantity", ""),
                item.get("unit", ""),
                _rubles(item.get("unit_price")),
                _rubles(item.get("line_total")),
            )
            for index, value in enumerate(values):
                _set_cell_text(
                    row.cells[index],
                    value,
                    bold=index == 5,
                    size=6.8,
                    align=(
                        WD_ALIGN_PARAGRAPH.RIGHT
                        if index in {2, 4, 5}
                        else WD_ALIGN_PARAGRAPH.CENTER
                        if index in {0, 3}
                        else None
                    ),
                )
                _set_cell_width(row.cells[index], breakdown_widths[index])
        total_row = breakdown.add_row()
        total_label = total_row.cells[0].merge(total_row.cells[4])
        _set_cell_text(
            total_label,
            "Итого по секции",
            bold=True,
            size=7.2,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        _set_cell_text(
            total_row.cells[5],
            f"{_rubles(line.get('document_line_total'))} ₽",
            bold=True,
            size=7.2,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        for cell in total_row.cells:
            _set_cell_shading(cell, HEADER_GRAY)
        _style_table(breakdown)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_glass_docx(context: dict) -> bytes:
    document = Document()
    _configure_document(document, landscape=False)
    _project_header(document, "Заказ стекла", context["project"])
    _add_project_document_warnings(document, context)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("КРОМКИ ПОЛИРОВАННЫЕ. Печать закалки не ставить")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(RED)

    table = document.add_table(rows=1, cols=8)
    headers = (
        "№",
        "Маркировка",
        "Стекло",
        "Ширина, мм",
        "Высота, мм",
        "Кол-во",
        "Площадь, м²",
        "Примечание",
    )
    for index, header in enumerate(headers):
        _set_cell_text(
            table.cell(0, index),
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row_data in context["glass_rows"]:
        row = table.add_row()
        values = (
            row_data["index"],
            row_data["marking"],
            row_data["glass_type"],
            format_dimension(row_data["width"]),
            format_dimension(row_data["height"]),
            row_data["qty"],
            f"{row_data['area']:.3f}",
            row_data["note"],
        )
        for index, value in enumerate(values):
            _set_cell_text(
                row.cells[index],
                value,
                size=7.5,
                align=WD_ALIGN_PARAGRAPH.CENTER if index != 2 else None,
            )
    total = table.add_row()
    total.cells[0].merge(total.cells[4])
    _set_cell_text(total.cells[0], "Итого", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(
        total.cells[5],
        context["glass_total_qty"],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        total.cells[6],
        f"{context['glass_total_area']:.3f}",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_table(table)

    quality = document.add_paragraph()
    quality.paragraph_format.space_before = Pt(8)
    quality_run = quality.add_run(
        "ОБРАЩАЮ ВНИМАНИЕ НА ПОВЫШЕННОЕ КАЧЕСТВО ИЗДЕЛИЙ\n"
        "Т.К. ФИРМА ИЗГОТАВЛИВАЕТ БЕЗРАМНОЕ ОСТЕКЛЕНИЕ\n"
        "ПЕРЕКОС ДИАГОНАЛИ НЕ ДОЛЖЕН ПРЕВЫШАТЬ 1-2ММ\n"
        "ГАБАРИТЫ СТЕКЛА НЕ ДОЛЖНЫ ПРЕВЫШАТЬ 2ММ"
    )
    quality_run.bold = True
    quality_run.font.name = "Arial"
    quality_run.font.size = Pt(10)
    quality_run.font.color.rgb = RGBColor.from_string(RED)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_paint_docx(context: dict) -> bytes:
    document = Document()
    _configure_document(document, landscape=False)
    pages = context["paint_pages"]
    if not pages:
        _project_header(document, "Заявка на покраску", context["project"])
        _add_project_document_warnings(document, context)
        document.add_paragraph("В проекте нет окрашиваемых профилей.")
    for page_index, page in enumerate(pages):
        if page_index:
            document.add_page_break()
        _project_header(document, "Заявка на покраску", context["project"])
        _add_project_document_warnings(document, context)
        color = document.add_paragraph()
        color.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = color.add_run(f"Цвет: {page['color']}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(RED)
        warning = document.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warning_run = warning.add_run("В СЧЕТЕ УКАЗАТЬ НОМЕР ЗАЯВКИ И ЦВЕТ ПРОФИЛЯ")
        warning_run.bold = True
        warning_run.font.name = "Arial"
        warning_run.font.size = Pt(9)
        warning_run.font.color.rgb = RGBColor.from_string(RED)

        table = document.add_table(rows=1, cols=6)
        headers = (
            "Артикул",
            "Сечение",
            "Кол-во",
            "Чистовые размеры",
            "С припуском 50 мм",
            "Общее, м.п.",
        )
        for index, header in enumerate(headers):
            _set_cell_text(
                table.cell(0, index),
                header,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        for group in page["groups"]:
            start = len(table.rows)
            for row_data in group["rows"]:
                row = table.add_row()
                _set_cell_text(
                    row.cells[0],
                    group["article"],
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                _set_cell_text(row.cells[1], group["name"], size=7.5)
                _set_cell_text(
                    row.cells[2], row_data["qty"], align=WD_ALIGN_PARAGRAPH.CENTER
                )
                _set_cell_text(
                    row.cells[3],
                    format_dimension(row_data["clean"]),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                _set_cell_text(
                    row.cells[4],
                    format_dimension(row_data["allowance"]),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                _set_cell_text(
                    row.cells[5],
                    f"{row_data['total_m']:.1f}".replace(".", ","),
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )
                row.height = Mm(11)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            end = len(table.rows) - 1
            article_cell = table.cell(start, 0)
            image_cell = table.cell(start, 1)
            if end > start:
                article_cell = article_cell.merge(table.cell(end, 0))
                image_cell = image_cell.merge(table.cell(end, 1))
            _set_cell_text(
                article_cell,
                group["article"],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(
                image_cell, group["name"], size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER
            )
            if group.get("note"):
                note = image_cell.add_paragraph(str(group["note"]))
                note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in note.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(RED)
                    run.font.size = Pt(8)
            source = image_stream(
                group.get("image"),
                group.get("image_data"),
                max_size=(700, 450),
            )
            _add_picture(image_cell, source, 22)

        total = table.add_row()
        total.cells[0].merge(total.cells[1])
        _set_cell_text(
            total.cells[0], "Итого", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT
        )
        _set_cell_text(
            total.cells[2],
            page["total_qty"],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text(
            total.cells[5],
            f"{page['total_m']:.1f}".replace(".", ","),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _style_table(table)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_hardware_order_docx(context: dict) -> bytes:
    document = Document()
    _configure_document(document, landscape=False)
    section = document.sections[0]
    section.top_margin = Mm(6)
    section.bottom_margin = Mm(6)
    pages = context["hardware_order_pages"]
    if not pages:
        pages = [{"system": "", "rows": [], "warning": ""}]

    for page_index, page in enumerate(pages):
        if page_index:
            document.add_page_break()

        dense_page = len(page["rows"]) > 24
        row_height_mm = 7 if dense_page else 8.2
        image_height_mm = 5.4 if dense_page else 6.8
        body_font_size = 6.2 if dense_page else 6.8
        compact_margin = 18 if dense_page else 45

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_after = Pt(0)
        run = heading.add_run(
            f"НАРЯД-ЗАКАЗ НА ФУРНИТУРУ — {getattr(context['project'], 'number', '')}"
        )
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)

        if page_index == 0:
            _add_project_document_warnings(document, context)

        system_heading = document.add_paragraph()
        system_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        system_heading.paragraph_format.space_after = Pt(1)
        system_run = system_heading.add_run(str(page["system"] or "БЕЗ СИСТЕМЫ"))
        system_run.bold = True
        system_run.font.name = "Arial"
        system_run.font.size = Pt(9)

        if page.get("warning"):
            warning = document.add_paragraph()
            warning.paragraph_format.space_after = Pt(1)
            warning_run = warning.add_run(str(page["warning"]))
            warning_run.bold = True
            warning_run.font.name = "Arial"
            warning_run.font.size = Pt(7)
            warning_run.font.color.rgb = RGBColor.from_string(RED)

        if page.get("note"):
            note = document.add_paragraph()
            note.paragraph_format.space_after = Pt(2)
            note_run = note.add_run(
                f"Примечание к комплектации: {page['note']}"
            )
            note_run.font.name = "Arial"
            note_run.font.size = Pt(7)
            note_run.italic = True

        table = document.add_table(rows=1, cols=8)
        table.autofit = False
        widths = (18, 23, 61, 22, 22, 13, 19, 16)
        headers = (
            "Артикул",
            "Эскиз",
            "Название",
            "Цвет",
            "Размер",
            "Этап",
            "Кол-во\n(общее в проекте)",
            "Единицы измерения",
        )
        for column, (header, width) in enumerate(zip(headers, widths)):
            cell = table.cell(0, column)
            _set_cell_width(cell, width)
            _set_cell_text(
                cell,
                header,
                bold=True,
                size=6.1 if dense_page else 6.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_margins(
                cell,
                top=compact_margin,
                start=50,
                bottom=compact_margin,
                end=50,
            )

        for row_data in page["rows"]:
            row = table.add_row()
            row.height = Mm(row_height_mm)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for column, width in enumerate(widths):
                _set_cell_width(row.cells[column], width)
            _set_cell_text(
                row.cells[0],
                row_data["article"],
                bold=True,
                size=6.1 if dense_page else 6.5,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(row.cells[1], "", size=body_font_size)
            _add_picture_fitted(
                row.cells[1],
                image_stream(row_data.get("image"), max_size=(700, 420)),
                max_width_mm=24,
                max_height_mm=image_height_mm,
            )
            _set_cell_text(row.cells[2], row_data["name"], size=body_font_size)
            _set_cell_text(
                row.cells[3],
                row_data.get("color") or "—",
                size=body_font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(
                row.cells[4],
                row_data.get("size") or "—",
                size=body_font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(
                row.cells[5],
                row_data["stage_text"],
                bold=True,
                size=body_font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(
                row.cells[6],
                row_data["qty_text"],
                bold=True,
                size=6.4 if dense_page else 7,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(
                row.cells[7],
                row_data["unit"],
                size=body_font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            for cell in row.cells:
                _set_cell_margins(
                    cell,
                    top=compact_margin,
                    start=50,
                    bottom=compact_margin,
                    end=50,
                )
        if not page["rows"]:
            row = table.add_row()
            cell = row.cells[0].merge(row.cells[7])
            _set_cell_text(
                cell,
                "Позиции не найдены",
                size=9,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        _style_table(table)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_project_docx(
    project: object,
    sections: Iterable[object],
    doc_type: str,
    quote: dict | None = None,
) -> bytes:
    context = build_project_document_context(project, sections, doc_type, quote=quote)
    if doc_type == "sketch":
        return _build_sketch_docx(context)
    if doc_type == "commercial" and quote is not None:
        return _build_commercial_docx(context)
    if doc_type == "glass":
        return _build_glass_docx(context)
    if doc_type == "paint":
        return _build_paint_docx(context)
    if doc_type == "hardware_order":
        return _build_hardware_order_docx(context)
    raise ValueError(
        "Word export is available only for sketch, commercial, glass, paint and hardware order documents"
    )
