import io
import zipfile
from types import SimpleNamespace

from docx import Document
from openpyxl import load_workbook

from engine.project_documents import build_project_document_context
from schemas import SectionCreate


def book_payload(**overrides):
    payload = {
        "name": "Секция КНИЖКА",
        "system": "КНИЖКА",
        "width": 3000,
        "height": 2500,
        "panels": 4,
        "quantity": 1,
        "book_system": "B25",
        "door_side": "both",
        "doors": 2,
        "book_left_door_hardware": "handle",
        "book_right_door_hardware": "lock",
        "book_left_door_opening": "inside_in",
        "book_right_door_opening": "outside_out",
        "book_left_stack_panels": 2,
        "book_obstacle_distance": 500,
        "book_handle_height": 1000,
        "compensator": "both",
    }
    payload.update(overrides)
    return payload


def test_book_calculation_has_authenticated_and_guest_endpoints(
    client,
    admin_headers,
):
    guest = client.post("/api/calculate/local/book", json=book_payload())
    assert guest.status_code == 200
    assert guest.json()["panels"][0]["glass_width_mm"] == 742.0
    assert guest.json()["normalized_config"]["door_layout"] == "both"

    assert client.post("/api/calculate/book", json=book_payload()).status_code == 403
    authenticated = client.post(
        "/api/calculate/book",
        headers=admin_headers,
        json=book_payload(),
    )
    assert authenticated.status_code == 200
    assert authenticated.json() == guest.json()


def test_existing_guest_calc_dispatches_to_book(client):
    response = client.post(
        "/api/projects/local/sections/calc",
        json={
            "project": {"number": "Гость", "customer": ""},
            "section": book_payload(),
        },
    )

    assert response.status_code == 200
    assert len(response.json()["panels"]) == 4
    assert response.json()["source_priority"][0] == "tz"


def test_book_api_returns_clear_calculation_error(client):
    response = client.post(
        "/api/calculate/local/book",
        json=book_payload(width=20),
    )

    assert response.status_code == 422
    assert "Суммарная ширина стекол" in response.json()["detail"]


def test_book_fields_save_copy_and_legacy_mapping(
    client,
    admin_headers,
    project,
):
    created = client.post(
        f"/api/projects/{project['id']}/sections",
        headers=admin_headers,
        json=book_payload(
            glass_supplied=False,
            book_left_door_width=710,
            book_right_door_width=720,
            book_left_fixed_left_enabled=True,
            book_left_fixed_left_width=450,
            book_left_fixed_right_enabled=True,
            book_left_fixed_right_width=460,
            book_right_fixed_left_enabled=True,
            book_right_fixed_left_width=470,
            book_right_fixed_right_enabled=True,
            book_right_fixed_right_width=480,
        ),
    )
    assert created.status_code == 201
    data = created.json()
    assert data["book_left_door_hardware"] == "handle"
    assert data["book_right_door_hardware"] == "lock"
    assert data["book_left_door_opening"] == "inside_in"
    assert data["book_right_door_opening"] == "outside_out"
    assert data["book_obstacle_distance"] == 500
    assert data["book_left_stack_panels"] == 2
    assert data["book_handle_height"] == 1000
    assert data["book_system"] == "B25"
    assert data["glass_supplied"] is False
    assert data["book_left_door_width"] == 710
    assert data["book_right_door_width"] == 720
    assert data["book_left_fixed_left_enabled"] is True
    assert data["book_left_fixed_left_width"] == 450
    assert data["book_left_fixed_right_enabled"] is True
    assert data["book_left_fixed_right_width"] == 460
    assert data["book_right_fixed_left_enabled"] is True
    assert data["book_right_fixed_left_width"] == 470
    assert data["book_right_fixed_right_enabled"] is True
    assert data["book_right_fixed_right_width"] == 480

    copied = client.post(
        f"/api/projects/{project['id']}/copy",
        headers=admin_headers,
    )
    assert copied.status_code == 201
    copied_book = next(
        row for row in copied.json()["sections"] if row["system"] == "КНИЖКА"
    )
    for field in (
        "book_left_door_hardware",
        "book_right_door_hardware",
        "book_left_door_opening",
        "book_right_door_opening",
        "book_obstacle_distance",
        "book_left_stack_panels",
        "book_handle_height",
        "book_system",
        "glass_supplied",
        "book_left_door_width",
        "book_right_door_width",
        "book_left_fixed_left_enabled",
        "book_left_fixed_left_width",
        "book_left_fixed_right_enabled",
        "book_left_fixed_right_width",
        "book_right_fixed_left_enabled",
        "book_right_fixed_left_width",
        "book_right_fixed_right_enabled",
        "book_right_fixed_right_width",
    ):
        assert copied_book[field] == data[field]
    client.delete(
        f"/api/projects/{copied.json()['id']}",
        headers=admin_headers,
    )

    legacy = client.post(
        f"/api/projects/{project['id']}/sections",
        headers=admin_headers,
        json=book_payload(
            name="Старая КНИЖКА",
            doors=1,
            door_side="Левая",
            door_type="Тип 4",
            door_opening="Наружу",
            book_system="Без каретки",
            book_left_door_hardware=None,
            book_right_door_hardware=None,
            book_left_door_opening=None,
            book_right_door_opening=None,
        ),
    )
    assert legacy.status_code == 201
    assert legacy.json()["door_side"] == "Левая"
    assert legacy.json()["book_left_door_hardware"] == "lock"
    assert legacy.json()["book_left_door_opening"] == "inside_out"
    assert legacy.json()["book_system"] == "B25"


def test_book_production_sheet_is_available_and_preliminary_configs_are_blocked(
    client,
    monkeypatch,
):
    confirmed = {
        "project": {"number": "Гость", "customer": ""},
        "section": book_payload(door_side="right", doors=1),
    }
    preview = client.post("/api/projects/local/sections/preview", json=confirmed)
    assert preview.status_code == 200
    assert "ПРЕДВАРИТЕЛЬНЫЙ ПРОИЗВОДСТВЕННЫЙ ЛИСТ" in preview.text
    assert "Сверловка D13/D6" in preview.text
    assert "RBP001" in preview.text
    assert "RBP002" in preview.text
    assert "RBP003" in preview.text
    assert "отверстие 93" not in preview.text

    docx = client.post("/api/projects/local/sections/docx", json=confirmed)
    xlsx = client.post("/api/projects/local/sections/xlsx", json=confirmed)
    assert docx.status_code == 200
    assert xlsx.status_code == 200
    assert docx.content.startswith(b"PK")
    assert xlsx.content.startswith(b"PK")

    word = Document(io.BytesIO(docx.content))
    word_text = "\n".join(
        paragraph.text for paragraph in word.paragraphs
    ) + "\n" + "\n".join(
        cell.text
        for table in word.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "ПРЕДВАРИТЕЛЬНЫЙ ПРОИЗВОДСТВЕННЫЙ ЛИСТ" in word_text
    assert "Сверловка D13/D6" in word_text
    assert all(article in word_text for article in ("RBP001", "RBP002", "RBP003"))
    assert "отверстие 93" not in word_text
    with zipfile.ZipFile(io.BytesIO(docx.content)) as archive:
        assert len(
            [name for name in archive.namelist() if name.startswith("word/media/")]
        ) >= 5

    workbook = load_workbook(io.BytesIO(xlsx.content), read_only=True, data_only=True)
    assert workbook.sheetnames == ["ПЛ КНИЖКА", "Комплектация"]
    excel_text = "\n".join(
        str(cell.value)
        for worksheet in workbook.worksheets
        for row in worksheet.iter_rows()
        for cell in row
        if cell.value is not None
    )
    assert "ПРЕДВАРИТЕЛЬНЫЙ ПРОИЗВОДСТВЕННЫЙ ЛИСТ" in excel_text
    assert "Сверловка D13/D6" in excel_text
    assert all(article in excel_text for article in ("RBP001", "RBP002", "RBP003"))
    assert "отверстие 93" not in excel_text
    with zipfile.ZipFile(io.BytesIO(xlsx.content)) as archive:
        assert len(
            [name for name in archive.namelist() if name.startswith("xl/media/")]
        ) >= 5

    monkeypatch.setattr("api.documents.generate_pdf", lambda html: b"%PDF-1.4\n%%EOF")
    pdf = client.post("/api/projects/local/sections/pdf", json=confirmed)
    assert pdf.status_code == 200
    assert pdf.content.startswith(b"%PDF")

    preliminary = {
        "project": {"number": "Гость", "customer": ""},
        "section": book_payload(angle_left=90),
    }
    blocked = client.post("/api/projects/local/sections/pdf", json=preliminary)
    assert blocked.status_code == 409
    assert "заблокированы" in blocked.json()["detail"]

    project_doc = client.post(
        "/api/projects/local/documents/glass/preview",
        json={
            "project": {"number": "Гость", "customer": ""},
            "sections": [book_payload()],
        },
    )
    assert project_doc.status_code == 501

    preliminary_project_doc = client.post(
        "/api/projects/local/documents/glass/preview",
        json={
            "project": {"number": "Гость", "customer": ""},
            "sections": [book_payload(angle_left=90)],
        },
    )
    assert preliminary_project_doc.status_code == 409

    confirmed_delivery = client.post(
        "/api/projects/local/documents/delivery/preview",
        json={
            "project": {"number": "Гость", "customer": ""},
            "sections": [book_payload()],
        },
    )
    assert confirmed_delivery.status_code == 200

    preliminary_delivery = client.post(
        "/api/projects/local/documents/delivery/preview",
        json={
            "project": {"number": "Гость", "customer": ""},
            "sections": [book_payload(angle_left=90)],
        },
    )
    assert preliminary_delivery.status_code == 409


def test_book_does_not_block_supported_parts_of_mixed_project_documents(
    client,
    monkeypatch,
):
    monkeypatch.setattr("api.documents.generate_pdf", lambda html: b"%PDF-1.4\n%%EOF")
    slide = {
        "name": "Секция СЛАЙД",
        "system": "СЛАЙД",
        "width": 2000,
        "height": 2400,
        "panels": 3,
        "quantity": 1,
        "rails": 3,
        "threshold": "Стандартный окраш",
        "painting_type": "RAL стандарт",
        "ral_color": "9016 МАТОВЫЙ",
        "first_panel_inside": "Справа",
    }
    payload = {
        "project": {"number": "MIXED-BOOK", "customer": "Тест"},
        "sections": [slide, book_payload(angle_left=90)],
    }
    project = SimpleNamespace(
        number="MIXED-BOOK",
        customer="Тест",
        glass_manual_rows="[]",
    )
    sections = [SectionCreate(**section) for section in payload["sections"]]

    for doc_type in ("glass", "paint"):
        responses = {
            extension: client.post(
                f"/api/projects/local/documents/{doc_type}/{extension}",
                json=payload,
            )
            for extension in ("preview", "pdf", "docx", "xlsx")
        }
        assert {response.status_code for response in responses.values()} == {200}
        preview = responses["preview"]
        assert "КНИЖКА не включена" in preview.text
        context = build_project_document_context(project, sections, doc_type)
        if doc_type == "glass":
            assert context["glass_rows"]
            expected_token = str(context["glass_rows"][0]["width"])
        else:
            paint_rows = [
                row for page in context["paint_pages"] for row in page["rows"]
            ]
            assert paint_rows
            expected_token = paint_rows[0]["article"]
        assert expected_token in preview.text


def test_book_hardware_order_contains_calculated_and_manual_rows(client):
    book = book_payload()
    response = client.post(
        "/api/projects/local/documents/hardware_order/preview",
        json={
            "project": {
                "number": "BOOK-HARDWARE",
                "customer": "Тест",
                "extra_components": (
                    '[{"sku":"BOOK-MANUAL","name":"Ручная позиция","qty":2}]'
                ),
            },
            "sections": [book],
        },
    )

    assert response.status_code == 200
    assert "Расчёт фурнитуры для системы КНИЖКА пока не реализован" not in response.text
    assert "RBP0004" in response.text
    assert ">1</td>" in response.text
    assert "BOOK-MANUAL" in response.text
